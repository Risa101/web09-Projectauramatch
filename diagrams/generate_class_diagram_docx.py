"""
AuraMatch — Conceptual Class Diagram (Word .docx)
Style: UML blue-box with <<stereotype>>, proper UML relationship arrows.
Each domain group on its own landscape A4 page.
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyArrowPatch, Polygon as MPoly
from matplotlib.lines import Line2D
import matplotlib.patheffects as pe
import numpy as np
import os, io, math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ─── UML Style Constants ───
FONT = "DejaVu Sans"
STEREO_FS = 10          # <<stereotype>> font
NAME_FS = 13            # class name font
ATTR_FS = 10.5          # attribute / method font
CLASS_W = 5.5           # box width
HEADER_H = 0.95         # header height (stereotype + name)
LINE_H = 0.36           # line height per attribute/method
PAD_L = 0.2             # left text padding

# Colors — pastel tones for readability
CLS_FILL = "#D6EEFF"    # very light blue fill
CLS_BORDER = "#2C2C2C"  # dark border
CLS_TEXT = "#000000"     # black text
METHOD_TEXT = "#000000"
REL_COLOR = "#2C2C2C"    # relationship line color
MULT_COLOR = "#C62828"   # multiplicity in red
LABEL_COLOR = "#C62828"  # annotation labels in red
BG_COLOR = "white"

OUT_DIR = "/Users/saridbutchuang/Desktop/web09_Projectauramatch/diagrams"

# ═══════════════════════════════════════════════
# GROUPS: each rendered on its own page
# stereotype: entity, boundary, control
# rel_type: association, aggregation, composition, dependency, generalization
# ═══════════════════════════════════════════════

GROUPS = [
    # ── Group 1: User & Auth ──
    {
        "title": "กลุ่มที่ 1: User & Authentication (ผู้ใช้งานและการยืนยันตัวตน)",
        "classes": {
            "User": {
                "stereotype": "entity",
                "attrs": ["-user_id : int", "-username : string", "-email : string",
                          "-password_hash : string", "-role : enum(admin, user)",
                          "-is_active : bool", "-created_at : datetime"],
                "methods": ["+register()", "+login()", "+logout()", "+getProfile()"],
                "pos": (1, 12),
            },
            "UserProfile": {
                "stereotype": "entity",
                "attrs": ["-profile_id : int", "-user_id : int [FK]",
                          "-first_name : string", "-last_name : string",
                          "-display_name : string", "-gender : enum",
                          "-avatar_url : string", "-birth_date : date"],
                "methods": ["+updateProfile()", "+uploadAvatar()"],
                "pos": (9, 12),
            },
            "PasswordReset": {
                "stereotype": "entity",
                "attrs": ["-reset_id : int", "-user_id : int [FK]",
                          "-token : string [unique]", "-expires_at : datetime",
                          "-created_at : datetime"],
                "methods": ["+resetPassword()", "+validateToken()"],
                "pos": (17, 12),
            },
        },
        "rels": [
            ("User", "UserProfile", "1", "1", "association", "right", "left"),
            ("User", "PasswordReset", "1", "*", "association", "right", "left"),
        ],
    },
    # ── Group 2: AI Analysis ──
    {
        "title": "กลุ่มที่ 2: AI Face Analysis & Color Science (การวิเคราะห์ใบหน้าและวิทยาศาสตร์สี)",
        "classes": {
            "AnalysisResult": {
                "stereotype": "entity",
                "attrs": ["-analysis_id : int", "-user_id : int [FK]",
                          "-image_path : string",
                          "-face_shape : enum(oval, round, square,\n     heart, oblong, diamond, triangle)",
                          "-skin_tone : enum(fair, light, medium,\n     tan, dark, deep)",
                          "-skin_undertone : enum(warm, cool, neutral)",
                          "-personal_color : enum(spring, summer,\n     autumn, winter)",
                          "-palette_id : int [FK]",
                          "-confidence_score : decimal",
                          "-created_at : datetime"],
                "methods": ["+analyzeFace(image)", "+getSkinLab()",
                            "+detectFaceShape()", "+getPersonalColor()"],
                "pos": (1, 14),
            },
            "ColorPalette": {
                "stereotype": "entity",
                "attrs": ["-palette_id : int", "-season : enum(4 seasons)",
                          "-sub_type : string (12 sub-seasons)",
                          "-best_colors : JSON {L*, a*, b*}",
                          "-avoid_colors : JSON {L*, a*, b*}",
                          "-description : text", "-makeup_tips : text"],
                "methods": ["+matchSeason(skin_lab)", "+getBestColors()",
                            "+getAvoidColors()"],
                "pos": (10, 14),
            },
            "AnalysisReview": {
                "stereotype": "entity",
                "attrs": ["-review_id : int", "-analysis_id : int [FK]",
                          "-user_id : int [FK]", "-is_accurate : bool",
                          "-comment : text", "-created_at : datetime"],
                "methods": ["+submitReview()"],
                "pos": (19, 14),
            },
        },
        "rels": [
            ("AnalysisResult", "ColorPalette", "*", "1", "aggregation", "right", "left"),
            ("AnalysisResult", "AnalysisReview", "1", "*", "association", "right", "left"),
        ],
    },
    # ── Group 3: Product Catalog ──
    {
        "title": "กลุ่มที่ 3: Product Catalog (แคตาล็อกสินค้า)",
        "classes": {
            "Product": {
                "stereotype": "entity",
                "attrs": ["-product_id : int", "-brand_id : int [FK]",
                          "-category_id : int [FK]", "-name : string",
                          "-description : text", "-price : decimal",
                          "-image_url : string", "-personal_color : string",
                          "-commission_rate : decimal", "-is_active : bool"],
                "methods": ["+getDetails()", "+getShades()", "+getLinks()"],
                "pos": (9, 16),
            },
            "Brand": {
                "stereotype": "entity",
                "attrs": ["-brand_id : int", "-name : string",
                          "-logo_url : string", "-website_url : string",
                          "-is_active : bool"],
                "methods": ["+getProducts()"],
                "pos": (1, 16),
            },
            "ProductCategory": {
                "stereotype": "entity",
                "attrs": ["-category_id : int", "-name : string",
                          "-description : text", "-icon_url : string"],
                "methods": ["+getProducts()"],
                "pos": (1, 10.5),
            },
            "ProductColorShade": {
                "stereotype": "entity",
                "attrs": ["-shade_id : int", "-product_id : int [FK]",
                          "-shade_name : string", "-shade_code : string",
                          "-hex_color : string", "-image_url : string"],
                "methods": ["+getColorLab() : {L,a,b}"],
                "pos": (17, 16),
            },
            "ProductLink": {
                "stereotype": "entity",
                "attrs": ["-link_id : int", "-product_id : int [FK]",
                          "-platform : enum(shopee, tiktok,\n     lazada, sephora, watsons)",
                          "-url : string", "-is_active : bool"],
                "methods": ["+trackClick()"],
                "pos": (17, 10.5),
            },
            "ProductTag": {
                "stereotype": "entity",
                "attrs": ["-tag_id : int", "-name : string [unique]"],
                "methods": [],
                "pos": (9, 10.5),
            },
            "ProductTagMap": {
                "stereotype": "entity",
                "attrs": ["-id : int", "-product_id : int [FK]",
                          "-tag_id : int [FK]"],
                "methods": [],
                "pos": (9, 7.5),
            },
        },
        "rels": [
            ("Brand", "Product", "1", "*", "association", "right", "left"),
            ("ProductCategory", "Product", "1", "*", "association", "bottom", "top"),
            ("Product", "ProductColorShade", "1", "*", "composition", "right", "left"),
            ("Product", "ProductLink", "1", "*", "aggregation", "right", "left"),
            ("Product", "ProductTagMap", "1", "*", "association", "bottom", "top"),
            ("ProductTag", "ProductTagMap", "1", "*", "association", "bottom", "top"),
        ],
    },
    # ── Group 4: Recommendation Engine ──
    {
        "title": "กลุ่มที่ 4: Recommendation Engine — Multi-Signal S1–S6 (เครื่องมือแนะนำสินค้า)",
        "classes": {
            "RecommendationRule": {
                "stereotype": "entity",
                "attrs": ["-rule_id : int", "-product_id : int [FK]",
                          "-face_shape : enum(7 + any)",
                          "-skin_tone : enum(6 + any)",
                          "-skin_undertone : enum(3 + any)",
                          "-personal_color : enum(4 + any)",
                          "-gender : enum", "-priority : int"],
                "methods": ["+matchRule(analysis)"],
                "pos": (1, 13),
            },
            "Recommendation": {
                "stereotype": "entity",
                "attrs": ["-recommendation_id : int",
                          "-analysis_id : int [FK]",
                          "-product_id : int [FK]",
                          "-score : decimal {S1+S2+S3+S4+S5+S6}"],
                "methods": ["+computeScore()", "+rank()"],
                "pos": (10, 13),
            },
            "RecommendationFeedback": {
                "stereotype": "entity",
                "attrs": ["-feedback_id : int",
                          "-recommendation_id : int [FK]",
                          "-user_id : int [FK]",
                          "-rating : enum(like, dislike)",
                          "-created_at : datetime"],
                "methods": ["+submitFeedback()"],
                "pos": (19, 13),
            },
            "Favorite": {
                "stereotype": "entity",
                "attrs": ["-favorite_id : int", "-user_id : int [FK]",
                          "-product_id : int [FK]",
                          "-created_at : datetime"],
                "methods": ["+toggle()"],
                "pos": (19, 8),
            },
        },
        "rels": [
            ("RecommendationRule", "Recommendation", "1..*", "1", "dependency", "right", "left"),
            ("Recommendation", "RecommendationFeedback", "1", "*", "association", "right", "left"),
        ],
    },
    # ── Group 5: Skin Concerns ──
    {
        "title": "กลุ่มที่ 5: Skin Concerns (ปัญหาผิวพรรณ)",
        "classes": {
            "SkinConcern": {
                "stereotype": "entity",
                "attrs": ["-concern_id : int", "-name : string [unique]",
                          "-description : text", "-icon_url : string"],
                "methods": [],
                "pos": (9, 11),
            },
            "UserSkinConcern": {
                "stereotype": "entity",
                "attrs": ["-id : int", "-user_id : int [FK]",
                          "-concern_id : int [FK]",
                          "-severity : enum(mild, moderate, severe)",
                          "-created_at : datetime"],
                "methods": [],
                "pos": (1, 11),
            },
            "ProductConcern": {
                "stereotype": "entity",
                "attrs": ["-id : int", "-product_id : int [FK]",
                          "-concern_id : int [FK]"],
                "methods": [],
                "pos": (17, 11),
            },
        },
        "rels": [
            ("SkinConcern", "UserSkinConcern", "1", "*", "association", "left", "right"),
            ("SkinConcern", "ProductConcern", "1", "*", "association", "right", "left"),
        ],
    },
    # ── Group 6: Gemini AI Chat ──
    {
        "title": "กลุ่มที่ 6: Gemini AI Chat (แชทบอท AI)",
        "classes": {
            "GeminiSession": {
                "stereotype": "entity",
                "attrs": ["-session_id : int", "-user_id : int [FK]",
                          "-analysis_id : int [FK] {optional}",
                          "-title : string", "-created_at : datetime"],
                "methods": ["+startChat()", "+getHistory()"],
                "pos": (1, 11),
            },
            "GeminiMessage": {
                "stereotype": "entity",
                "attrs": ["-message_id : int", "-session_id : int [FK]",
                          "-role : enum(user, model)",
                          "-prompt : text", "-response : text",
                          "-image_input : string", "-image_output : string",
                          "-created_at : datetime"],
                "methods": ["+send()", "+getResponse()"],
                "pos": (11, 11),
            },
            "GeminiService": {
                "stereotype": "control",
                "attrs": [],
                "methods": ["+generateResponse(prompt)",
                            "+analyzeImage(image)",
                            "+buildContext(session)"],
                "pos": (6, 6),
            },
        },
        "rels": [
            ("GeminiSession", "GeminiMessage", "1", "*", "composition", "right", "left"),
            ("GeminiService", "GeminiSession", "", "", "dependency", "top", "bottom"),
        ],
    },
    # ── Group 7: Photo Editor ──
    {
        "title": "กลุ่มที่ 7: Photo Editor (ตัวแต่งรูปภาพ)",
        "classes": {
            "PhotoEdit": {
                "stereotype": "entity",
                "attrs": ["-edit_id : int", "-user_id : int [FK]",
                          "-original_image : string",
                          "-edited_image : string",
                          "-edit_config : JSON",
                          "-source : enum(upload, analysis, gemini)",
                          "-created_at : datetime"],
                "methods": ["+applyFilter()", "+save()"],
                "pos": (1, 12),
            },
            "SavedLook": {
                "stereotype": "entity",
                "attrs": ["-look_id : int", "-user_id : int [FK]",
                          "-name : string",
                          "-category : string(day, night, party, work)",
                          "-makeup_data : JSON", "-filter_data : JSON",
                          "-thumbnail_url : string"],
                "methods": ["+saveLook()", "+loadLook()"],
                "pos": (10, 12),
            },
            "EditFilter": {
                "stereotype": "boundary",
                "attrs": ["-filter_id : int", "-name : string",
                          "-category : enum(beauty, color, light,\n     vintage, makeup)",
                          "-config : JSON", "-is_active : bool"],
                "methods": [],
                "pos": (1, 6.5),
            },
            "EditSticker": {
                "stereotype": "boundary",
                "attrs": ["-sticker_id : int", "-name : string",
                          "-category : enum(face, decoration,\n     text, emoji)",
                          "-image_url : string", "-is_active : bool"],
                "methods": [],
                "pos": (10, 6.5),
            },
        },
        "rels": [
            ("PhotoEdit", "EditFilter", "", "", "dependency", "bottom", "top"),
            ("PhotoEdit", "EditSticker", "", "", "dependency", "bottom", "top"),
        ],
    },
    # ── Group 8: Blog, Reviews & Affiliate ──
    {
        "title": "กลุ่มที่ 8: Blog, Reviews & Affiliate Tracking (บล็อก, รีวิว และการติดตาม)",
        "classes": {
            "BlogCategory": {
                "stereotype": "entity",
                "attrs": ["-category_id : int", "-name : string",
                          "-description : text"],
                "methods": [],
                "pos": (1, 14),
            },
            "BlogPost": {
                "stereotype": "entity",
                "attrs": ["-post_id : int", "-author_id : int [FK]",
                          "-category_id : int [FK]",
                          "-title : string", "-slug : string [unique]",
                          "-content : text", "-views : int",
                          "-is_published : bool"],
                "methods": ["+publish()", "+incrementViews()"],
                "pos": (9, 14),
            },
            "ProductReview": {
                "stereotype": "entity",
                "attrs": ["-review_id : int", "-product_id : int [FK]",
                          "-user_id : int [FK]", "-rating : int(1-5)",
                          "-comment : text", "-image_url : string",
                          "-platform : enum", "-is_verified : bool"],
                "methods": ["+submitReview()"],
                "pos": (17, 14),
            },
            "ClickLog": {
                "stereotype": "entity",
                "attrs": ["-log_id : int", "-link_id : int [FK]",
                          "-user_id : int [FK]",
                          "-platform : enum", "-ip_address : string",
                          "-user_agent : text", "-clicked_at : datetime"],
                "methods": ["+logClick()"],
                "pos": (1, 8.5),
            },
            "Commission": {
                "stereotype": "entity",
                "attrs": ["-commission_id : int", "-link_id : int [FK]",
                          "-user_id : int [FK]", "-platform : enum",
                          "-amount : decimal",
                          "-status : enum(pending, confirmed, paid)",
                          "-clicked_at : datetime"],
                "methods": ["+confirm()", "+pay()"],
                "pos": (9, 8.5),
            },
        },
        "rels": [
            ("BlogCategory", "BlogPost", "1", "*", "association", "right", "left"),
            ("ClickLog", "Commission", "", "", "dependency", "right", "left"),
        ],
    },
    # ── Group 9: Admin & Analytics ──
    {
        "title": "กลุ่มที่ 9: Admin, Notifications & Analytics (แอดมิน, การแจ้งเตือน และการวิเคราะห์)",
        "classes": {
            "Notification": {
                "stereotype": "entity",
                "attrs": ["-notification_id : int", "-user_id : int [FK]",
                          "-title : string", "-message : text",
                          "-type : enum(recommendation,\n     promotion, system, review)",
                          "-is_read : bool", "-created_at : datetime"],
                "methods": ["+markRead()", "+send()"],
                "pos": (1, 13),
            },
            "AdminLog": {
                "stereotype": "entity",
                "attrs": ["-log_id : int", "-admin_id : int [FK]",
                          "-action : string", "-table_name : string",
                          "-record_id : int",
                          "-old_value : JSON", "-new_value : JSON",
                          "-ip_address : string", "-created_at : datetime"],
                "methods": ["+logAction()"],
                "pos": (10, 13),
            },
            "UserBehavior": {
                "stereotype": "entity",
                "attrs": ["-id : int", "-user_id : int [FK]",
                          "-session_id : string",
                          "-event_type : string (product_view,\n     search, filter, click)",
                          "-event_data : JSON", "-page : string",
                          "-created_at : datetime"],
                "methods": ["+trackEvent()"],
                "pos": (19, 13),
            },
            "SearchHistory": {
                "stereotype": "entity",
                "attrs": ["-search_id : int", "-user_id : int [FK]",
                          "-keyword : string", "-created_at : datetime"],
                "methods": [],
                "pos": (1, 7.5),
            },
            "Banner": {
                "stereotype": "boundary",
                "attrs": ["-banner_id : int", "-title : string",
                          "-image_url : string", "-link_url : string",
                          "-position : enum(home_top,\n     home_middle, sidebar)",
                          "-starts_at : datetime", "-ends_at : datetime",
                          "-is_active : bool"],
                "methods": [],
                "pos": (10, 7.5),
            },
        },
        "rels": [],
    },
    # ── Group 10: Overall Relationships ──
    {
        "title": "กลุ่มที่ 10: ภาพรวมความสัมพันธ์ — Core Entities",
        "classes": {
            "User": {
                "stereotype": "entity",
                "attrs": ["-user_id : int",
                          "-username : string",
                          "-email : string",
                          "-role : enum"],
                "methods": ["+login()", "+register()"],
                "pos": (1, 16),
            },
            "AnalysisResult": {
                "stereotype": "entity",
                "attrs": ["-analysis_id : int",
                          "-face_shape : enum",
                          "-personal_color : enum"],
                "methods": ["+analyzeFace()"],
                "pos": (9, 16),
            },
            "ColorPalette": {
                "stereotype": "entity",
                "attrs": ["-palette_id : int",
                          "-season : enum",
                          "-best_colors : JSON"],
                "methods": ["+matchSeason()"],
                "pos": (17, 16),
            },
            "Recommendation": {
                "stereotype": "entity",
                "attrs": ["-recommendation_id : int",
                          "-score : decimal"],
                "methods": ["+computeScore()"],
                "pos": (9, 10.5),
            },
            "Product": {
                "stereotype": "entity",
                "attrs": ["-product_id : int",
                          "-name : string",
                          "-price : decimal"],
                "methods": ["+getDetails()"],
                "pos": (17, 10.5),
            },
            "GeminiSession": {
                "stereotype": "entity",
                "attrs": ["-session_id : int",
                          "-title : string"],
                "methods": ["+startChat()"],
                "pos": (1, 10.5),
            },
            "Favorite": {
                "stereotype": "entity",
                "attrs": ["-favorite_id : int"],
                "methods": ["+toggle()"],
                "pos": (9, 5.5),
            },
            "PhotoEdit": {
                "stereotype": "entity",
                "attrs": ["-edit_id : int",
                          "-source : enum"],
                "methods": ["+save()"],
                "pos": (1, 5.5),
            },
        },
        "rels": [
            ("User", "AnalysisResult", "1", "*", "association", "right", "left"),
            ("AnalysisResult", "ColorPalette", "*", "1", "aggregation", "right", "left"),
            ("AnalysisResult", "Recommendation", "1", "*", "association", "bottom", "top"),
            ("Product", "Recommendation", "1", "*", "association", "left", "right"),
            ("User", "GeminiSession", "1", "*", "association", "bottom", "top"),
            ("User", "Favorite", "1", "*", "association", "bottom", "top"),
            ("Product", "Favorite", "1", "*", "association", "bottom", "top"),
            ("User", "PhotoEdit", "1", "*", "association", "bottom", "top"),
        ],
    },
]


# ═══════════════════════════════════════════════
# Drawing engine
# ═══════════════════════════════════════════════

def _hex(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16)/255 for i in (0,2,4))

STEREO_COLORS = {
    "entity": _hex(CLS_FILL),        # very light blue
    "boundary": _hex("#D7F5D7"),     # very light green
    "control": _hex("#FFE8CC"),      # very light orange
}

def _count_display_lines(lines):
    """Count actual display lines (considering \\n in strings)."""
    count = 0
    for line in lines:
        count += line.count("\n") + 1
    return count

def _box_h(attrs, methods):
    n_a = max(_count_display_lines(attrs), 1)
    n_m = _count_display_lines(methods)
    h = HEADER_H + n_a * LINE_H + 0.3
    if n_m:
        h += n_m * LINE_H + 0.18
    else:
        h += 0.1
    return h


def _draw_class(ax, name, data):
    x, y_top = data["pos"]
    attrs = data["attrs"]
    methods = data["methods"]
    stereo = data.get("stereotype", "entity")
    w = CLASS_W
    h = _box_h(attrs, methods)
    fill = STEREO_COLORS.get(stereo, STEREO_COLORS["entity"])

    # Main box
    rect = Rectangle((x, y_top - h), w, h,
                      facecolor=fill, edgecolor=_hex(CLS_BORDER),
                      linewidth=2.0, zorder=2)
    ax.add_patch(rect)

    # Header: stereotype + name
    ax.text(x + w/2, y_top - 0.25, f"\u00ab{stereo}\u00bb",
            ha="center", va="center", fontsize=STEREO_FS,
            fontfamily=FONT, color=_hex(CLS_TEXT), fontstyle="italic", zorder=4)
    ax.text(x + w/2, y_top - 0.65, name,
            ha="center", va="center", fontsize=NAME_FS, fontweight="bold",
            fontfamily=FONT, color=_hex(CLS_TEXT), zorder=4)

    # Header separator
    sep_y = y_top - HEADER_H
    ax.plot([x, x + w], [sep_y, sep_y],
            color=_hex(CLS_BORDER), linewidth=1.5, zorder=3)

    # Attributes
    cy = sep_y - 0.28
    for a in attrs:
        sub_lines = a.split("\n")
        for sl in sub_lines:
            ax.text(x + PAD_L, cy, sl, ha="left", va="center",
                    fontsize=ATTR_FS, fontfamily=FONT, color=_hex(CLS_TEXT), zorder=4)
            cy -= LINE_H

    # Method separator + methods
    if methods:
        msep_y = cy + LINE_H * 0.25
        ax.plot([x, x + w], [msep_y, msep_y],
                color=_hex(CLS_BORDER), linewidth=1.0, zorder=3)
        cy -= 0.08
        for m in methods:
            sub_lines = m.split("\n")
            for sl in sub_lines:
                ax.text(x + PAD_L, cy, sl, ha="left", va="center",
                        fontsize=ATTR_FS, fontfamily=FONT, color=_hex(METHOD_TEXT), zorder=4)
                cy -= LINE_H

    data["_bbox"] = (x, y_top - h, x + w, y_top)
    data["_h"] = h


def _anchor(bbox, side):
    x0, y0, x1, y1 = bbox
    cx, cy = (x0+x1)/2, (y0+y1)/2
    return {"top":(cx,y1), "bottom":(cx,y0), "left":(x0,cy), "right":(x1,cy)}[side]


def _draw_diamond(ax, pos, side, filled=False):
    """Draw a small diamond (aggregation/composition) at connection point."""
    x, y = pos
    sz = 0.22
    if side in ("left", "right"):
        dx = -sz if side == "left" else sz
        pts = [(x, y), (x + dx, y + sz*0.6), (x + 2*dx, y), (x + dx, y - sz*0.6)]
        new_pos = (x + 2*dx, y)
    else:
        dy = -sz if side == "bottom" else sz
        pts = [(x, y), (x + sz*0.6, y + dy), (x, y + 2*dy), (x - sz*0.6, y + dy)]
        new_pos = (x, y + 2*dy)

    diamond = MPoly(pts, closed=True,
                    facecolor=_hex(CLS_BORDER) if filled else "white",
                    edgecolor=_hex(CLS_BORDER), linewidth=1.5, zorder=6)
    ax.add_patch(diamond)
    return new_pos


def _draw_rel(ax, cls_dict, ca, cb, ma, mb, rel_type, side_a, side_b):
    ba = cls_dict[ca]["_bbox"]
    bb = cls_dict[cb]["_bbox"]
    pa = list(_anchor(ba, side_a))
    pb = list(_anchor(bb, side_b))

    # Draw endpoint decorations first
    actual_pa = pa[:]
    actual_pb = pb[:]

    if rel_type == "aggregation":
        actual_pa = list(_draw_diamond(ax, pa, side_a, filled=False))
    elif rel_type == "composition":
        actual_pa = list(_draw_diamond(ax, pa, side_a, filled=True))

    # Line style
    ls = "--" if rel_type == "dependency" else "-"
    lw = 1.5

    # Draw orthogonal line
    if side_a in ("top","bottom") and side_b in ("top","bottom"):
        mid_y = (actual_pa[1]+actual_pb[1])/2
        ax.plot([actual_pa[0],actual_pa[0],actual_pb[0],actual_pb[0]],
                [actual_pa[1],mid_y,mid_y,actual_pb[1]],
                color=_hex(REL_COLOR), linewidth=lw, linestyle=ls, zorder=1)
    elif side_a in ("left","right") and side_b in ("left","right"):
        mid_x = (actual_pa[0]+actual_pb[0])/2
        ax.plot([actual_pa[0],mid_x,mid_x,actual_pb[0]],
                [actual_pa[1],actual_pa[1],actual_pb[1],actual_pb[1]],
                color=_hex(REL_COLOR), linewidth=lw, linestyle=ls, zorder=1)
    else:
        if side_a in ("top","bottom"):
            ax.plot([actual_pa[0],actual_pa[0],actual_pb[0]],
                    [actual_pa[1],actual_pb[1],actual_pb[1]],
                    color=_hex(REL_COLOR), linewidth=lw, linestyle=ls, zorder=1)
        else:
            ax.plot([actual_pa[0],actual_pb[0],actual_pb[0]],
                    [actual_pa[1],actual_pa[1],actual_pb[1]],
                    color=_hex(REL_COLOR), linewidth=lw, linestyle=ls, zorder=1)

    # Arrow head for dependency
    if rel_type == "dependency":
        _draw_arrowhead(ax, actual_pb, side_b)

    # Multiplicity labels
    d = 0.35
    def _mpos(p, side):
        if side == "top": return (p[0]+0.2, p[1]+d)
        if side == "bottom": return (p[0]+0.2, p[1]-d)
        if side == "right": return (p[0]+d, p[1]+0.22)
        return (p[0]-d-0.1, p[1]+0.22)

    if ma:
        ta = _mpos(pa, side_a)
        ax.text(ta[0], ta[1], ma, fontsize=11, fontweight="bold",
                color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)
    if mb:
        tb = _mpos(pb, side_b)
        ax.text(tb[0], tb[1], mb, fontsize=11, fontweight="bold",
                color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)


def _draw_arrowhead(ax, pos, side):
    """Draw open arrowhead (dependency)."""
    x, y = pos
    sz = 0.2
    if side == "left":
        pts = [(x-sz, y+sz*0.6), (x, y), (x-sz, y-sz*0.6)]
    elif side == "right":
        pts = [(x+sz, y+sz*0.6), (x, y), (x+sz, y-sz*0.6)]
    elif side == "top":
        pts = [(x-sz*0.6, y+sz), (x, y), (x+sz*0.6, y+sz)]
    else:
        pts = [(x-sz*0.6, y-sz), (x, y), (x+sz*0.6, y-sz)]

    ax.plot([p[0] for p in pts], [p[1] for p in pts],
            color=_hex(REL_COLOR), linewidth=1.5, zorder=6)


def render_group(group):
    """Render one group to PNG bytes."""
    cls = group["classes"]
    rels = group["rels"]

    # Auto-calculate bounds
    all_x = [d["pos"][0] for d in cls.values()]
    all_y = [d["pos"][1] for d in cls.values()]
    min_x = min(all_x) - 1.5
    max_x = max(all_x) + CLASS_W + 2
    min_y = min(all_y) - 7
    max_y = max(all_y) + 2

    fw = max(max_x - min_x, 14)
    fh = max(max_y - min_y, 8)

    fig, ax = plt.subplots(figsize=(fw * 1.0, fh * 1.0))
    ax.set_xlim(min_x, max_x)
    ax.set_ylim(min_y, max_y)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")

    for name, data in cls.items():
        _draw_class(ax, name, data)

    for r in rels:
        _draw_rel(ax, cls, *r)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════
# Build Word document
# ═══════════════════════════════════════════════

doc = Document()

# Page setup: A4 landscape
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ── Cover page ──
for _ in range(4):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AuraMatch")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Conceptual Class Diagram")
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph("")

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run("แพลตฟอร์มแนะนำเครื่องสำอางด้วย AI\nAI-Powered Beauty & Cosmetics Recommendation Platform")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("35 Entities  ·  37 Relationships  ·  10 Domain Groups")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

doc.add_paragraph("")

# Legend
legend_title = doc.add_paragraph()
legend_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = legend_title.add_run("สัญลักษณ์ (Legend)")
r.font.size = Pt(14)
r.font.bold = True

legend_items = [
    "■ สีฟ้า = <<entity>> — Entity Class (ตารางฐานข้อมูล)",
    "■ สีเขียว = <<boundary>> — Boundary Class (ส่วนติดต่อผู้ใช้)",
    "■ สีส้ม = <<control>> — Control Class (ตัวควบคุมธุรกิจ)",
    "── เส้นทึบ = Association (ความสัมพันธ์)",
    "- - เส้นประ = Dependency (การพึ่งพา)",
    "◇── = Aggregation (การรวมตัว — ส่วนประกอบแยกได้)",
    "◆── = Composition (การประกอบ — ส่วนประกอบแยกไม่ได้)",
    "1, * = Multiplicity (จำนวนความสัมพันธ์)",
]
for item in legend_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

# ── Table of Contents ──
doc.add_page_break()
toc = doc.add_paragraph()
toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc.add_run("สารบัญ — Table of Contents")
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("")

for i, g in enumerate(GROUPS, 1):
    entity_names = ", ".join(g["classes"].keys())
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    r = p.add_run(f"หน้า {i+2}  —  {g['title']}")
    r.font.size = Pt(12)
    r.font.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(3)
    r2 = p2.add_run(f"Entities: {entity_names}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

# ── Relationship summary table ──
doc.add_page_break()
rel_title = doc.add_paragraph()
rel_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = rel_title.add_run("ตารางสรุปความสัมพันธ์ทั้งหมด — All Relationships")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("")

table = doc.add_table(rows=1, cols=5, style="Light Grid Accent 1")
headers = ["#", "From Entity", "To Entity", "Multiplicity", "Relationship Type"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for rr in p.runs:
            rr.font.bold = True
            rr.font.size = Pt(9)

all_rels = [
    ("User", "UserProfile", "1 — 1", "Association"),
    ("User", "PasswordReset", "1 — *", "Association"),
    ("User", "AnalysisResult", "1 — *", "Association"),
    ("AnalysisResult", "ColorPalette", "* — 1", "Aggregation"),
    ("AnalysisResult", "AnalysisReview", "1 — *", "Association"),
    ("AnalysisResult", "Recommendation", "1 — *", "Association"),
    ("Product", "Recommendation", "1 — *", "Association"),
    ("Recommendation", "RecommendationFeedback", "1 — *", "Association"),
    ("User", "RecommendationFeedback", "1 — *", "Association"),
    ("Product", "RecommendationRule", "1 — *", "Association"),
    ("Brand", "Product", "1 — *", "Association"),
    ("ProductCategory", "Product", "1 — *", "Association"),
    ("Product", "ProductColorShade", "1 — *", "Composition"),
    ("Product", "ProductLink", "1 — *", "Aggregation"),
    ("Product", "ProductTagMap", "1 — *", "Association"),
    ("ProductTag", "ProductTagMap", "1 — *", "Association"),
    ("User", "Favorite", "1 — *", "Association"),
    ("Product", "Favorite", "1 — *", "Association"),
    ("User", "ProductReview", "1 — *", "Association"),
    ("Product", "ProductReview", "1 — *", "Association"),
    ("SkinConcern", "UserSkinConcern", "1 — *", "Association"),
    ("SkinConcern", "ProductConcern", "1 — *", "Association"),
    ("User", "UserSkinConcern", "1 — *", "Association"),
    ("Product", "ProductConcern", "1 — *", "Association"),
    ("User", "GeminiSession", "1 — *", "Association"),
    ("GeminiSession", "GeminiMessage", "1 — *", "Composition"),
    ("GeminiSession", "AnalysisResult", "* — 0..1", "Association"),
    ("User", "PhotoEdit", "1 — *", "Association"),
    ("User", "SavedLook", "1 — *", "Association"),
    ("PhotoEdit", "EditFilter", "— —", "Dependency"),
    ("PhotoEdit", "EditSticker", "— —", "Dependency"),
    ("User", "BlogPost", "1 — *", "Association"),
    ("BlogCategory", "BlogPost", "1 — *", "Association"),
    ("ProductLink", "ClickLog", "1 — *", "Association"),
    ("ProductLink", "Commission", "1 — *", "Association"),
    ("User", "Notification", "1 — *", "Association"),
    ("User", "SearchHistory", "1 — *", "Association"),
    ("User", "AdminLog", "1 — *", "Association"),
    ("User", "UserBehavior", "1 — *", "Association"),
]
for idx, (fr, to, mult, rtype) in enumerate(all_rels, 1):
    row = table.add_row().cells
    row[0].text = str(idx)
    row[1].text = fr
    row[2].text = to
    row[3].text = mult
    row[4].text = rtype
    for cell in row:
        for p in cell.paragraphs:
            for rr in p.runs:
                rr.font.size = Pt(9)

# ── Per-group diagram pages ──
for group in GROUPS:
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(group["title"])
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # Render and insert diagram
    img_buf = render_group(group)
    doc.add_picture(img_buf, width=Cm(25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    entity_list = ", ".join(group["classes"].keys())
    r = cap.add_run(f"Entities ({len(group['classes'])}): {entity_list}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    r.font.italic = True

# ── Save ──
out_path = os.path.join(OUT_DIR, "AuraMatch_Class_Diagram.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
