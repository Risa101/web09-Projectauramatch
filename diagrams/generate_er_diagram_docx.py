"""
AuraMatch — ER Diagram (Word .docx)
การแปลง ORM Class เป็นตารางในฐานข้อมูลเชิงสัมพันธ์
แต่ละกลุ่มโดเมนอยู่คนละหน้า landscape A4
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch
import io, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ─── Style Constants ───
FONT = "DejaVu Sans"
TBL_NAME_FS = 13
COL_FS = 9.5
TBL_W = 5.0
LINE_H = 0.32
HEADER_H = 0.7
PAD_L = 0.18

PK_COLOR = "#1565C0"
FK_COLOR = "#C62828"
COL_COLOR = "#212121"
TBL_FILL = "#E3F2FD"
TBL_BORDER = "#1565C0"
REF_FILL = "#FFF8E1"
REF_BORDER = "#F9A825"
REL_COLOR = "#424242"
MULT_COLOR = "#C62828"
BG = "white"

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def _hex(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16)/255 for i in (0, 2, 4))


# ═══════════════════════════════════════
# ER GROUPS — each group = one page
# columns: list of (name, type, flags)
#   flags: "PK", "FK", "UK", "NN"
# ═══════════════════════════════════════

GROUPS = [
    # ── Group 1 ──
    {
        "title": "กลุ่มที่ 1: User & Authentication (ผู้ใช้งานและการยืนยันตัวตน)",
        "tables": {
            "users": {
                "class": "User",
                "file": "models/user.py",
                "pos": (1, 14),
                "cols": [
                    ("user_id", "INT", "PK,AI"),
                    ("username", "VARCHAR(50)", "UK,NN"),
                    ("email", "VARCHAR(100)", "UK,NN"),
                    ("password_hash", "VARCHAR(255)", "NN"),
                    ("role", "ENUM('admin','user')", ""),
                    ("is_active", "INT default 1", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "user_profiles": {
                "class": "UserProfile",
                "file": "models/misc.py",
                "pos": (8, 14),
                "cols": [
                    ("profile_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,UK,NN"),
                    ("first_name", "VARCHAR(100)", ""),
                    ("last_name", "VARCHAR(100)", ""),
                    ("display_name", "VARCHAR(100)", ""),
                    ("avatar_url", "VARCHAR(255)", ""),
                    ("birth_date", "DATE", ""),
                    ("gender", "ENUM(4 values)", ""),
                    ("nationality", "VARCHAR(100)", ""),
                    ("bio", "TEXT", ""),
                    ("updated_at", "DATETIME", ""),
                ],
            },
            "password_resets": {
                "class": "PasswordReset",
                "file": "models/misc.py",
                "pos": (15, 14),
                "cols": [
                    ("reset_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("token", "VARCHAR(255)", "UK,NN"),
                    ("expires_at", "DATETIME", "NN"),
                    ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("users", "user_profiles", "1", "1", "right", "left"),
            ("users", "password_resets", "1", "*", "right", "left"),
        ],
    },
    # ── Group 2 ──
    {
        "title": "กลุ่มที่ 2: Product Catalog (แคตาล็อกสินค้า)",
        "tables": {
            "brands": {
                "class": "Brand",
                "file": "models/product.py",
                "pos": (1, 16),
                "cols": [
                    ("brand_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("logo_url", "VARCHAR(255)", ""),
                    ("website_url", "VARCHAR(255)", ""),
                    ("description", "TEXT", ""),
                    ("is_active", "INT default 1", ""),
                ],
            },
            "product_categories": {
                "class": "ProductCategory",
                "file": "models/product.py",
                "pos": (1, 10),
                "cols": [
                    ("category_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""),
                    ("icon_url", "VARCHAR(255)", ""),
                ],
            },
            "products": {
                "class": "Product",
                "file": "models/product.py",
                "pos": (8, 16),
                "cols": [
                    ("product_id", "INT", "PK,AI"),
                    ("brand_id", "INT → brands", "FK"),
                    ("category_id", "INT → categories", "FK"),
                    ("name", "VARCHAR(200)", "NN"),
                    ("description", "TEXT", ""),
                    ("price", "DECIMAL(10,2)", ""),
                    ("image_url", "VARCHAR(255)", ""),
                    ("personal_color", "VARCHAR(50)", ""),
                    ("commission_rate", "DECIMAL(5,2)", ""),
                    ("is_active", "INT default 1", ""),
                ],
            },
            "product_links": {
                "class": "ProductLink",
                "file": "models/product.py",
                "pos": (15, 16),
                "cols": [
                    ("link_id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("platform", "ENUM(6 platforms)", "NN"),
                    ("url", "VARCHAR(500)", "NN"),
                    ("is_active", "INT default 1", ""),
                ],
            },
            "product_color_shades": {
                "class": "ProductColorShade",
                "file": "models/product.py",
                "pos": (15, 10),
                "cols": [
                    ("shade_id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("shade_name", "VARCHAR(100)", ""),
                    ("hex_color", "VARCHAR(7)", ""),
                    ("image_url", "VARCHAR(255)", ""),
                ],
            },
            "product_tags": {
                "class": "ProductTag",
                "file": "models/product.py",
                "pos": (8, 10),
                "cols": [
                    ("tag_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "UK,NN"),
                ],
            },
            "product_tag_map": {
                "class": "ProductTagMap",
                "file": "models/product.py",
                "pos": (8, 7),
                "cols": [
                    ("id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("tag_id", "INT → product_tags", "FK,NN"),
                ],
            },
        },
        "rels": [
            ("brands", "products", "1", "*", "right", "left"),
            ("product_categories", "products", "1", "*", "bottom", "top"),
            ("products", "product_links", "1", "*", "right", "left"),
            ("products", "product_color_shades", "1", "*", "right", "left"),
            ("products", "product_tag_map", "1", "*", "bottom", "top"),
            ("product_tags", "product_tag_map", "1", "*", "bottom", "top"),
        ],
    },
    # ── Group 3 ──
    {
        "title": "กลุ่มที่ 3: AI Analysis & Color Science (การวิเคราะห์ใบหน้าและวิทยาศาสตร์สี)",
        "tables": {
            "color_palettes": {
                "class": "ColorPalette",
                "file": "models/analysis.py",
                "pos": (1, 14),
                "cols": [
                    ("palette_id", "INT", "PK,AI"),
                    ("season", "ENUM(4 seasons)", "NN"),
                    ("sub_type", "VARCHAR(50)", ""),
                    ("description", "TEXT", ""),
                    ("best_colors", "JSON {L*,a*,b*}", ""),
                    ("avoid_colors", "JSON {L*,a*,b*}", ""),
                    ("makeup_tips", "TEXT", ""),
                ],
            },
            "analysis_results": {
                "class": "AnalysisResult",
                "file": "models/analysis.py",
                "pos": (8, 14),
                "cols": [
                    ("analysis_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("image_path", "VARCHAR(255)", ""),
                    ("gender", "ENUM(3 values)", ""),
                    ("ethnicity", "ENUM(8 values)", ""),
                    ("face_shape", "ENUM(7 shapes)", ""),
                    ("skin_tone", "ENUM(6 tones)", ""),
                    ("skin_undertone", "ENUM(3 values)", ""),
                    ("personal_color", "ENUM(4 seasons)", ""),
                    ("palette_id", "INT → color_palettes", "FK"),
                    ("confidence_score", "DECIMAL(5,2)", ""),
                ],
            },
            "analysis_reviews": {
                "class": "AnalysisReview",
                "file": "models/analysis.py",
                "pos": (16, 14),
                "cols": [
                    ("review_id", "INT", "PK,AI"),
                    ("analysis_id", "INT → analysis_results", "FK,NN"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("is_accurate", "INT", ""),
                    ("comment", "TEXT", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("color_palettes", "analysis_results", "1", "*", "right", "left"),
            ("analysis_results", "analysis_reviews", "1", "*", "right", "left"),
        ],
    },
    # ── Group 4 ──
    {
        "title": "กลุ่มที่ 4: Recommendation Engine — Multi-Signal S1–S6 (เครื่องมือแนะนำสินค้า)",
        "tables": {
            "recommendation_rules": {
                "class": "RecommendationRule",
                "file": "models/recommendation.py",
                "pos": (1, 14),
                "cols": [
                    ("rule_id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("face_shape", "ENUM(7+any)", ""),
                    ("skin_tone", "ENUM(6+any)", ""),
                    ("skin_undertone", "ENUM(3+any)", ""),
                    ("personal_color", "ENUM(4+any)", ""),
                    ("gender", "ENUM(2+any)", ""),
                    ("priority", "INT default 0", ""),
                ],
            },
            "recommendations": {
                "class": "Recommendation",
                "file": "models/recommendation.py",
                "pos": (8, 14),
                "cols": [
                    ("recommendation_id", "INT", "PK,AI"),
                    ("analysis_id", "INT → analysis_results", "FK,NN"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("score", "DECIMAL(5,2)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "recommendation_feedback": {
                "class": "RecommendationFeedback",
                "file": "models/recommendation.py",
                "pos": (15, 14),
                "cols": [
                    ("feedback_id", "INT", "PK,AI"),
                    ("recommendation_id", "INT → recommendations", "FK,NN"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("rating", "ENUM('like','dislike')", "NN"),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "favorites": {
                "class": "Favorite",
                "file": "models/recommendation.py",
                "pos": (15, 8.5),
                "cols": [
                    ("favorite_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("recommendation_rules", "recommendations", "1..*", "1", "right", "left"),
            ("recommendations", "recommendation_feedback", "1", "*", "right", "left"),
        ],
    },
    # ── Group 5 ──
    {
        "title": "กลุ่มที่ 5: Skin Concerns (ปัญหาผิวพรรณ)",
        "tables": {
            "skin_concerns": {
                "class": "SkinConcern",
                "file": "models/misc.py",
                "pos": (8, 12),
                "cols": [
                    ("concern_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""),
                    ("icon_url", "VARCHAR(255)", ""),
                ],
            },
            "user_skin_concerns": {
                "class": "UserSkinConcern",
                "file": "models/misc.py",
                "pos": (1, 12),
                "cols": [
                    ("id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("concern_id", "INT → skin_concerns", "FK,NN"),
                    ("severity", "ENUM(3 levels)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "product_concerns": {
                "class": "ProductConcern",
                "file": "models/misc.py",
                "pos": (15, 12),
                "cols": [
                    ("id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("concern_id", "INT → skin_concerns", "FK,NN"),
                ],
            },
        },
        "rels": [
            ("skin_concerns", "user_skin_concerns", "1", "*", "left", "right"),
            ("skin_concerns", "product_concerns", "1", "*", "right", "left"),
        ],
    },
    # ── Group 6 ──
    {
        "title": "กลุ่มที่ 6: Gemini AI Chat (แชทบอท AI)",
        "tables": {
            "gemini_sessions": {
                "class": "GeminiSession",
                "file": "models/gemini.py",
                "pos": (1, 12),
                "cols": [
                    ("session_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("analysis_id", "INT → analysis_results", "FK"),
                    ("title", "VARCHAR(200)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "gemini_messages": {
                "class": "GeminiMessage",
                "file": "models/gemini.py",
                "pos": (9, 12),
                "cols": [
                    ("message_id", "INT", "PK,AI"),
                    ("session_id", "INT → gemini_sessions", "FK,NN"),
                    ("role", "ENUM('user','model')", "NN"),
                    ("prompt", "TEXT", ""),
                    ("response", "TEXT", ""),
                    ("image_input", "VARCHAR(255)", ""),
                    ("image_output", "VARCHAR(255)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("gemini_sessions", "gemini_messages", "1", "*", "right", "left"),
        ],
    },
    # ── Group 7 ──
    {
        "title": "กลุ่มที่ 7: Photo Editor (ตัวแต่งรูปภาพ)",
        "tables": {
            "photo_edits": {
                "class": "PhotoEdit",
                "file": "models/photo_editor.py",
                "pos": (1, 13),
                "cols": [
                    ("edit_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("original_image", "VARCHAR(255)", ""),
                    ("edited_image", "VARCHAR(255)", ""),
                    ("edit_config", "JSON", ""),
                    ("source", "ENUM('upload','analysis','gemini')", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "saved_looks": {
                "class": "SavedLook",
                "file": "models/saved_look.py",
                "pos": (8, 13),
                "cols": [
                    ("look_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("category", "VARCHAR(50)", ""),
                    ("makeup_data", "JSON", "NN"),
                    ("filter_data", "JSON", ""),
                    ("thumbnail_url", "VARCHAR(500)", ""),
                ],
            },
            "edit_filters": {
                "class": "EditFilter",
                "file": "models/photo_editor.py",
                "pos": (15, 13),
                "cols": [
                    ("filter_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("category", "ENUM(5 types)", "NN"),
                    ("thumbnail_url", "VARCHAR(255)", ""),
                    ("config", "JSON", ""),
                    ("is_active", "INT default 1", ""),
                ],
            },
            "edit_stickers": {
                "class": "EditSticker",
                "file": "models/photo_editor.py",
                "pos": (15, 8),
                "cols": [
                    ("sticker_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("category", "ENUM(4 types)", "NN"),
                    ("image_url", "VARCHAR(255)", ""),
                    ("is_active", "INT default 1", ""),
                ],
            },
        },
        "rels": [],
    },
    # ── Group 8 ──
    {
        "title": "กลุ่มที่ 8: Blog & Reviews & Affiliate (บล็อก, รีวิว และ Affiliate)",
        "tables": {
            "blog_categories": {
                "class": "BlogCategory",
                "file": "models/blog.py",
                "pos": (1, 14),
                "cols": [
                    ("category_id", "INT", "PK,AI"),
                    ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""),
                ],
            },
            "blog_posts": {
                "class": "BlogPost",
                "file": "models/blog.py",
                "pos": (8, 14),
                "cols": [
                    ("post_id", "INT", "PK,AI"),
                    ("category_id", "INT → blog_categories", "FK"),
                    ("author_id", "INT → users", "FK,NN"),
                    ("title", "VARCHAR(200)", "NN"),
                    ("slug", "VARCHAR(200)", "UK"),
                    ("content", "TEXT", ""),
                    ("views", "INT default 0", ""),
                    ("is_published", "INT default 0", ""),
                ],
            },
            "product_reviews": {
                "class": "ProductReview",
                "file": "models/misc.py",
                "pos": (15, 14),
                "cols": [
                    ("review_id", "INT", "PK,AI"),
                    ("product_id", "INT → products", "FK,NN"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("rating", "INT", "NN"),
                    ("comment", "TEXT", ""),
                    ("platform", "ENUM(4 values)", ""),
                    ("is_verified", "INT default 0", ""),
                ],
            },
            "click_logs": {
                "class": "ClickLog",
                "file": "models/commission.py",
                "pos": (1, 8.5),
                "cols": [
                    ("log_id", "INT", "PK,AI"),
                    ("link_id", "INT → product_links", "FK,NN"),
                    ("user_id", "INT → users", "FK"),
                    ("platform", "ENUM(3 platforms)", "NN"),
                    ("ip_address", "VARCHAR(45)", ""),
                    ("clicked_at", "DATETIME", ""),
                ],
            },
            "commissions": {
                "class": "Commission",
                "file": "models/commission.py",
                "pos": (8, 8.5),
                "cols": [
                    ("commission_id", "INT", "PK,AI"),
                    ("link_id", "INT → product_links", "FK,NN"),
                    ("user_id", "INT → users", "FK"),
                    ("platform", "ENUM(3 platforms)", "NN"),
                    ("amount", "DECIMAL(10,2)", ""),
                    ("status", "ENUM(3 statuses)", ""),
                ],
            },
        },
        "rels": [
            ("blog_categories", "blog_posts", "1", "*", "right", "left"),
            ("click_logs", "commissions", "", "", "right", "left"),
        ],
    },
    # ── Group 9 ──
    {
        "title": "กลุ่มที่ 9: Admin, Notifications & Analytics (แอดมิน, แจ้งเตือน, วิเคราะห์พฤติกรรม)",
        "tables": {
            "notifications": {
                "class": "Notification",
                "file": "models/misc.py",
                "pos": (1, 13),
                "cols": [
                    ("notification_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK,NN"),
                    ("title", "VARCHAR(200)", "NN"),
                    ("message", "TEXT", ""),
                    ("type", "ENUM(4 types)", ""),
                    ("is_read", "INT default 0", ""),
                ],
            },
            "admin_logs": {
                "class": "AdminLog",
                "file": "models/misc.py",
                "pos": (8, 13),
                "cols": [
                    ("log_id", "INT", "PK,AI"),
                    ("admin_id", "INT → users", "FK,NN"),
                    ("action", "VARCHAR(100)", "NN"),
                    ("table_name", "VARCHAR(100)", ""),
                    ("record_id", "INT", ""),
                    ("old_value", "JSON", ""),
                    ("new_value", "JSON", ""),
                ],
            },
            "user_behaviors": {
                "class": "UserBehavior",
                "file": "models/behavior.py",
                "pos": (15, 13),
                "cols": [
                    ("id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK"),
                    ("session_id", "VARCHAR(100)", ""),
                    ("event_type", "VARCHAR(50)", "NN"),
                    ("event_data", "JSON", ""),
                    ("page", "VARCHAR(50)", ""),
                ],
            },
            "search_history": {
                "class": "SearchHistory",
                "file": "models/misc.py",
                "pos": (1, 7.5),
                "cols": [
                    ("search_id", "INT", "PK,AI"),
                    ("user_id", "INT → users", "FK"),
                    ("keyword", "VARCHAR(255)", "NN"),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "banners": {
                "class": "Banner",
                "file": "models/misc.py",
                "pos": (8, 7.5),
                "cols": [
                    ("banner_id", "INT", "PK,AI"),
                    ("title", "VARCHAR(200)", ""),
                    ("image_url", "VARCHAR(255)", "NN"),
                    ("link_url", "VARCHAR(500)", ""),
                    ("position", "ENUM(3 positions)", ""),
                    ("is_active", "INT default 1", ""),
                ],
            },
        },
        "rels": [],
    },
    # ── Group 10: Overview ──
    {
        "title": "กลุ่มที่ 10: ภาพรวมความสัมพันธ์ — Core Entities (Overview)",
        "tables": {
            "users": {
                "class": "User", "file": "models/user.py",
                "pos": (1, 16),
                "cols": [("user_id", "INT", "PK"), ("username", "VARCHAR(50)", "UK")],
            },
            "analysis_results": {
                "class": "AnalysisResult", "file": "models/analysis.py",
                "pos": (8, 16),
                "cols": [("analysis_id", "INT", "PK"), ("user_id", "INT → users", "FK"),
                         ("face_shape", "ENUM", ""), ("personal_color", "ENUM", "")],
            },
            "color_palettes": {
                "class": "ColorPalette", "file": "models/analysis.py",
                "pos": (15, 16),
                "cols": [("palette_id", "INT", "PK"), ("season", "ENUM", ""),
                         ("best_colors", "JSON", "")],
            },
            "recommendations": {
                "class": "Recommendation", "file": "models/recommendation.py",
                "pos": (8, 10.5),
                "cols": [("recommendation_id", "INT", "PK"), ("score", "DECIMAL", "")],
            },
            "products": {
                "class": "Product", "file": "models/product.py",
                "pos": (15, 10.5),
                "cols": [("product_id", "INT", "PK"), ("name", "VARCHAR(200)", ""),
                         ("price", "DECIMAL", "")],
            },
            "gemini_sessions": {
                "class": "GeminiSession", "file": "models/gemini.py",
                "pos": (1, 10.5),
                "cols": [("session_id", "INT", "PK"), ("title", "VARCHAR(200)", "")],
            },
            "favorites": {
                "class": "Favorite", "file": "models/recommendation.py",
                "pos": (8, 5.5),
                "cols": [("favorite_id", "INT", "PK")],
            },
            "photo_edits": {
                "class": "PhotoEdit", "file": "models/photo_editor.py",
                "pos": (1, 5.5),
                "cols": [("edit_id", "INT", "PK"), ("source", "ENUM", "")],
            },
        },
        "rels": [
            ("users", "analysis_results", "1", "*", "right", "left"),
            ("analysis_results", "color_palettes", "*", "1", "right", "left"),
            ("analysis_results", "recommendations", "1", "*", "bottom", "top"),
            ("products", "recommendations", "1", "*", "left", "right"),
            ("users", "gemini_sessions", "1", "*", "bottom", "top"),
            ("users", "favorites", "1", "*", "bottom", "top"),
            ("products", "favorites", "1", "*", "bottom", "top"),
            ("users", "photo_edits", "1", "*", "bottom", "top"),
        ],
    },
]


# ═══════════════════════════════════════
# Drawing engine
# ═══════════════════════════════════════

def _box_h(cols):
    n = max(len(cols), 1)
    return HEADER_H + n * LINE_H + 0.35


def _draw_table(ax, tbl_name, data):
    x, y_top = data["pos"]
    cols = data["cols"]
    cls_name = data["class"]
    w = TBL_W
    h = _box_h(cols)
    fill = _hex(TBL_FILL)
    border = _hex(TBL_BORDER)

    # Box
    rect = FancyBboxPatch((x, y_top - h), w, h,
                           boxstyle="round,pad=0.05",
                           facecolor=fill, edgecolor=border,
                           linewidth=2.0, zorder=2)
    ax.add_patch(rect)

    # Header: table name + class name
    ax.text(x + w/2, y_top - 0.22, tbl_name,
            ha="center", va="center", fontsize=TBL_NAME_FS, fontweight="bold",
            fontfamily=FONT, color=_hex(TBL_BORDER), zorder=4)
    ax.text(x + w/2, y_top - 0.52, f"({cls_name})",
            ha="center", va="center", fontsize=COL_FS, fontstyle="italic",
            fontfamily=FONT, color=_hex("#616161"), zorder=4)

    # Header separator
    sep_y = y_top - HEADER_H
    ax.plot([x + 0.08, x + w - 0.08], [sep_y, sep_y],
            color=border, linewidth=1.2, zorder=3)

    # Columns
    cy = sep_y - 0.25
    for col_name, col_type, flags in cols:
        flag_set = set(f.strip() for f in flags.split(",")) if flags else set()

        # Icon
        if "PK" in flag_set:
            icon = "PK "
            name_color = _hex(PK_COLOR)
        elif "FK" in flag_set:
            icon = "FK "
            name_color = _hex(FK_COLOR)
        else:
            icon = "     "
            name_color = _hex(COL_COLOR)

        # Build text
        badge = ""
        if "PK" in flag_set:
            badge = " [PK]"
        if "FK" in flag_set:
            badge += " [FK]"
        if "UK" in flag_set:
            badge += " [UK]"
        if "NN" in flag_set:
            badge += " [NN]"

        txt = f"{icon}{col_name} : {col_type}{badge}"
        ax.text(x + PAD_L, cy, txt, ha="left", va="center",
                fontsize=COL_FS, fontfamily=FONT, color=name_color, zorder=4)
        cy -= LINE_H

    data["_bbox"] = (x, y_top - h, x + w, y_top)


def _anchor(bbox, side):
    x0, y0, x1, y1 = bbox
    cx, cy = (x0+x1)/2, (y0+y1)/2
    return {"top": (cx, y1), "bottom": (cx, y0),
            "left": (x0, cy), "right": (x1, cy)}[side]


def _draw_crow_foot(ax, pos, side):
    """Draw crow's foot (many) at endpoint."""
    x, y = pos
    sz = 0.18
    if side == "left":
        ax.plot([x-sz, x, x-sz], [y+sz*0.7, y, y-sz*0.7],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=6)
    elif side == "right":
        ax.plot([x+sz, x, x+sz], [y+sz*0.7, y, y-sz*0.7],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=6)
    elif side == "top":
        ax.plot([x-sz*0.7, x, x+sz*0.7], [y+sz, y, y+sz],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=6)
    else:
        ax.plot([x-sz*0.7, x, x+sz*0.7], [y-sz, y, y-sz],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=6)


def _draw_rel(ax, tbl_dict, ta, tb, ma, mb, side_a, side_b):
    ba = tbl_dict[ta]["_bbox"]
    bb = tbl_dict[tb]["_bbox"]
    pa = list(_anchor(ba, side_a))
    pb = list(_anchor(bb, side_b))

    # Draw orthogonal line
    if side_a in ("top", "bottom") and side_b in ("top", "bottom"):
        mid_y = (pa[1]+pb[1])/2
        ax.plot([pa[0], pa[0], pb[0], pb[0]], [pa[1], mid_y, mid_y, pb[1]],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=1)
    elif side_a in ("left", "right") and side_b in ("left", "right"):
        mid_x = (pa[0]+pb[0])/2
        ax.plot([pa[0], mid_x, mid_x, pb[0]], [pa[1], pa[1], pb[1], pb[1]],
                color=_hex(REL_COLOR), linewidth=1.5, zorder=1)
    else:
        if side_a in ("top", "bottom"):
            ax.plot([pa[0], pa[0], pb[0]], [pa[1], pb[1], pb[1]],
                    color=_hex(REL_COLOR), linewidth=1.5, zorder=1)
        else:
            ax.plot([pa[0], pb[0], pb[0]], [pa[1], pa[1], pb[1]],
                    color=_hex(REL_COLOR), linewidth=1.5, zorder=1)

    # Crow's foot for "many" side
    if mb == "*":
        _draw_crow_foot(ax, pb, side_b)

    # Multiplicity labels
    d = 0.35
    def _mpos(p, side):
        if side == "top": return (p[0]+0.2, p[1]+d)
        if side == "bottom": return (p[0]+0.2, p[1]-d)
        if side == "right": return (p[0]+d, p[1]+0.22)
        return (p[0]-d-0.1, p[1]+0.22)

    if ma:
        ta_pos = _mpos(pa, side_a)
        ax.text(ta_pos[0], ta_pos[1], ma, fontsize=11, fontweight="bold",
                color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)
    if mb:
        tb_pos = _mpos(pb, side_b)
        ax.text(tb_pos[0], tb_pos[1], mb, fontsize=11, fontweight="bold",
                color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)


def render_group(group):
    tbls = group["tables"]
    rels = group["rels"]

    all_x = [d["pos"][0] for d in tbls.values()]
    all_y = [d["pos"][1] for d in tbls.values()]
    min_x = min(all_x) - 1.5
    max_x = max(all_x) + TBL_W + 2
    min_y = min(all_y) - max(_box_h(d["cols"]) for d in tbls.values()) - 1.5
    max_y = max(all_y) + 2

    fw = max(max_x - min_x, 12)
    fh = max(max_y - min_y, 7)

    fig, ax = plt.subplots(figsize=(fw, fh))
    ax.set_xlim(min_x, max_x)
    ax.set_ylim(min_y, max_y)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")

    for name, data in tbls.items():
        _draw_table(ax, name, data)

    for r in rels:
        _draw_rel(ax, tbls, *r)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════
# Build Word document
# ═══════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ── Cover page ──
for _ in range(4):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AuraMatch")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Entity-Relationship Diagram")
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph("")

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run("การแปลง ORM Class เป็นตารางในฐานข้อมูลเชิงสัมพันธ์\nClass-to-Relational Table Mapping")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("23 Tables  ·  39 Relationships  ·  10 Domain Groups  ·  MySQL (InnoDB)")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

doc.add_paragraph("")

# Legend
legend_title = doc.add_paragraph()
legend_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = legend_title.add_run("สัญลักษณ์ (Legend)")
r.font.size = Pt(14)
r.font.bold = True

legend_items = [
    "PK = Primary Key (คีย์หลัก)",
    "FK = Foreign Key (คีย์นอก)",
    "UK = Unique Key (คีย์ไม่ซ้ำ)",
    "NN = Not Null (ห้ามว่าง)",
    "AI = Auto Increment",
    "── เส้นทึบ = Relationship (ความสัมพันธ์)",
    "1, * = Multiplicity (จำนวน: หนึ่ง, หลาย)",
    "< = Crow's Foot (ฝั่ง Many)",
]
for item in legend_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

# ── Class-to-Table Mapping summary ──
doc.add_page_break()

map_title = doc.add_paragraph()
map_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = map_title.add_run("ตารางสรุปการแปลง ORM Class → Relational Table")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("")

table = doc.add_table(rows=1, cols=5, style="Light Grid Accent 1")
headers = ["#", "ORM Class (Python)", "Table Name (MySQL)", "Source File", "Domain Group"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for rr in p.runs:
            rr.font.bold = True
            rr.font.size = Pt(9)

idx = 0
for gi, g in enumerate(GROUPS):
    for tbl_name, tbl_data in g["tables"].items():
        # Skip ref tables in overview group
        if gi == len(GROUPS) - 1:
            continue
        idx += 1
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = tbl_data["class"]
        row[2].text = tbl_name
        row[3].text = tbl_data["file"]
        row[4].text = g["title"].split("(")[0].strip()
        for cell in row:
            for p in cell.paragraphs:
                for rr in p.runs:
                    rr.font.size = Pt(9)

# ── Relationship summary ──
doc.add_page_break()

rel_title = doc.add_paragraph()
rel_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = rel_title.add_run("ตารางสรุปความสัมพันธ์ทั้งหมด (All Relationships)")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("")

rtbl = doc.add_table(rows=1, cols=5, style="Light Grid Accent 1")
rh = ["#", "Parent Table (1)", "Child Table (N)", "FK Column", "ON DELETE"]
for i, h in enumerate(rh):
    rtbl.rows[0].cells[i].text = h
    for p in rtbl.rows[0].cells[i].paragraphs:
        for rr in p.runs:
            rr.font.bold = True
            rr.font.size = Pt(9)

all_fk = [
    ("users", "user_profiles", "user_id", "CASCADE"),
    ("users", "password_resets", "user_id", "CASCADE"),
    ("users", "analysis_results", "user_id", "CASCADE"),
    ("users", "analysis_reviews", "user_id", "CASCADE"),
    ("users", "favorites", "user_id", "CASCADE"),
    ("users", "notifications", "user_id", "CASCADE"),
    ("users", "gemini_sessions", "user_id", "CASCADE"),
    ("users", "blog_posts", "author_id", "—"),
    ("users", "photo_edits", "user_id", "CASCADE"),
    ("users", "saved_looks", "user_id", "CASCADE"),
    ("users", "product_reviews", "user_id", "CASCADE"),
    ("users", "recommendation_feedback", "user_id", "CASCADE"),
    ("users", "search_history", "user_id", "SET NULL"),
    ("users", "user_skin_concerns", "user_id", "CASCADE"),
    ("users", "admin_logs", "admin_id", "—"),
    ("users", "user_behaviors", "user_id", "SET NULL"),
    ("users", "click_logs", "user_id", "SET NULL"),
    ("users", "commissions", "user_id", "SET NULL"),
    ("brands", "products", "brand_id", "SET NULL"),
    ("product_categories", "products", "category_id", "SET NULL"),
    ("products", "product_links", "product_id", "CASCADE"),
    ("products", "product_color_shades", "product_id", "CASCADE"),
    ("products", "product_tag_map", "product_id", "CASCADE"),
    ("product_tags", "product_tag_map", "tag_id", "CASCADE"),
    ("products", "product_reviews", "product_id", "CASCADE"),
    ("products", "favorites", "product_id", "CASCADE"),
    ("products", "recommendation_rules", "product_id", "CASCADE"),
    ("products", "recommendations", "product_id", "—"),
    ("products", "product_concerns", "product_id", "CASCADE"),
    ("color_palettes", "analysis_results", "palette_id", "—"),
    ("analysis_results", "recommendations", "analysis_id", "CASCADE"),
    ("analysis_results", "analysis_reviews", "analysis_id", "CASCADE"),
    ("analysis_results", "gemini_sessions", "analysis_id", "SET NULL"),
    ("recommendations", "recommendation_feedback", "recommendation_id", "CASCADE"),
    ("product_links", "click_logs", "link_id", "—"),
    ("product_links", "commissions", "link_id", "—"),
    ("blog_categories", "blog_posts", "category_id", "SET NULL"),
    ("skin_concerns", "user_skin_concerns", "concern_id", "—"),
    ("skin_concerns", "product_concerns", "concern_id", "—"),
    ("gemini_sessions", "gemini_messages", "session_id", "CASCADE"),
]
for i, (parent, child, fk_col, on_del) in enumerate(all_fk, 1):
    row = rtbl.add_row().cells
    row[0].text = str(i)
    row[1].text = parent
    row[2].text = child
    row[3].text = fk_col
    row[4].text = on_del
    for cell in row:
        for p in cell.paragraphs:
            for rr in p.runs:
                rr.font.size = Pt(9)

# ── Per-group diagram pages ──
for group in GROUPS:
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(group["title"])
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    img_buf = render_group(group)
    doc.add_picture(img_buf, width=Cm(25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_list = ", ".join(f"{n} ({d['class']})" for n, d in group["tables"].items())
    r = cap.add_run(f"Tables ({len(group['tables'])}): {tbl_list}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    r.font.italic = True

# ── Save ──
out_path = os.path.join(OUT_DIR, "AuraMatch_ER_Diagram.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
