"""
AuraMatch — รายงานการแปลงคลาสเป็นตารางในฐานข้อมูลเชิงสัมพันธ์
พร้อม ER Diagram ประกอบ
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import io, os, textwrap
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Style ───
FONT = "DejaVu Sans"
TBL_NAME_FS = 13
COL_FS = 9.5
TBL_W = 5.0
LINE_H = 0.32
HEADER_H = 0.7
PAD_L = 0.18

PK_COLOR = "#1565C0"
FK_COLOR = "#C62828"
COL_COLOR = "#212121"
TBL_FILL = "#E3F2FD"
TBL_BORDER = "#1565C0"
REL_COLOR = "#424242"
MULT_COLOR = "#C62828"


def _hex(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16)/255 for i in (0, 2, 4))


# ═══════════════════════════════════════
# ER GROUPS for diagrams
# ═══════════════════════════════════════

ER_GROUPS = [
    {
        "title": "กลุ่มที่ 1: User & Authentication",
        "tables": {
            "users": {
                "class": "User", "file": "models/user.py", "pos": (1, 14),
                "cols": [
                    ("user_id", "INT", "PK,AI"), ("username", "VARCHAR(50)", "UK,NN"),
                    ("email", "VARCHAR(100)", "UK,NN"), ("password_hash", "VARCHAR(255)", "NN"),
                    ("role", "ENUM('admin','user')", ""), ("is_active", "INT", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "user_profiles": {
                "class": "UserProfile", "file": "models/misc.py", "pos": (8, 14),
                "cols": [
                    ("profile_id", "INT", "PK,AI"), ("user_id", "INT", "FK,UK,NN"),
                    ("first_name", "VARCHAR(100)", ""), ("last_name", "VARCHAR(100)", ""),
                    ("display_name", "VARCHAR(100)", ""), ("avatar_url", "VARCHAR(255)", ""),
                    ("birth_date", "DATE", ""), ("gender", "ENUM(4 values)", ""),
                    ("nationality", "VARCHAR(100)", ""), ("bio", "TEXT", ""),
                    ("updated_at", "DATETIME", ""),
                ],
            },
            "password_resets": {
                "class": "PasswordReset", "file": "models/misc.py", "pos": (15, 14),
                "cols": [
                    ("reset_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("token", "VARCHAR(255)", "UK,NN"), ("expires_at", "DATETIME", "NN"),
                    ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("users", "user_profiles", "1", "1", "right", "left"),
            ("users", "password_resets", "1", "*", "right", "left"),
        ],
    },
    {
        "title": "กลุ่มที่ 2: Product Catalog",
        "tables": {
            "brands": {
                "class": "Brand", "file": "models/product.py", "pos": (1, 16),
                "cols": [
                    ("brand_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("logo_url", "VARCHAR(255)", ""), ("website_url", "VARCHAR(255)", ""),
                    ("description", "TEXT", ""), ("is_active", "INT", ""),
                ],
            },
            "product_categories": {
                "class": "ProductCategory", "file": "models/product.py", "pos": (1, 10),
                "cols": [
                    ("category_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""), ("icon_url", "VARCHAR(255)", ""),
                ],
            },
            "products": {
                "class": "Product", "file": "models/product.py", "pos": (8, 16),
                "cols": [
                    ("product_id", "INT", "PK,AI"), ("brand_id", "INT", "FK"),
                    ("category_id", "INT", "FK"), ("name", "VARCHAR(200)", "NN"),
                    ("description", "TEXT", ""), ("price", "DECIMAL(10,2)", ""),
                    ("image_url", "VARCHAR(255)", ""), ("personal_color", "VARCHAR(50)", ""),
                    ("commission_rate", "DECIMAL(5,2)", ""), ("is_active", "INT", ""),
                ],
            },
            "product_links": {
                "class": "ProductLink", "file": "models/product.py", "pos": (15, 16),
                "cols": [
                    ("link_id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("platform", "ENUM(6 values)", "NN"), ("url", "VARCHAR(500)", "NN"),
                    ("is_active", "INT", ""),
                ],
            },
            "product_color_shades": {
                "class": "ProductColorShade", "file": "models/product.py", "pos": (15, 10),
                "cols": [
                    ("shade_id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("shade_name", "VARCHAR(100)", ""), ("shade_code", "VARCHAR(20)", ""),
                    ("hex_color", "VARCHAR(7)", ""),
                    ("image_url", "VARCHAR(255)", ""), ("is_active", "INT", ""),
                ],
            },
            "product_tags": {
                "class": "ProductTag", "file": "models/product.py", "pos": (8, 10),
                "cols": [("tag_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "UK,NN")],
            },
            "product_tag_map": {
                "class": "ProductTagMap", "file": "models/product.py", "pos": (8, 7),
                "cols": [
                    ("id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("tag_id", "INT", "FK,NN"),
                ],
            },
        },
        "rels": [
            ("brands", "products", "1", "*", "right", "left"),
            ("product_categories", "products", "1", "*", "bottom", "top"),
            ("products", "product_links", "1", "*", "right", "left"),
            ("products", "product_color_shades", "1", "*", "right", "left"),
            ("products", "product_tag_map", "1", "*", "bottom", "top"),
            ("product_tags", "product_tag_map", "1", "*", "bottom", "top"),
        ],
    },
    {
        "title": "กลุ่มที่ 3: AI Analysis & Color Science",
        "tables": {
            "color_palettes": {
                "class": "ColorPalette", "file": "models/analysis.py", "pos": (1, 14),
                "cols": [
                    ("palette_id", "INT", "PK,AI"), ("season", "ENUM(4 seasons)", "NN"),
                    ("sub_type", "VARCHAR(50)", ""), ("description", "TEXT", ""),
                    ("best_colors", "JSON", ""), ("avoid_colors", "JSON", ""),
                    ("makeup_tips", "TEXT", ""),
                ],
            },
            "analysis_results": {
                "class": "AnalysisResult", "file": "models/analysis.py", "pos": (8, 14),
                "cols": [
                    ("analysis_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("image_path", "VARCHAR(255)", ""),
                    ("gender", "ENUM(3 values)", ""),
                    ("ethnicity", "ENUM(8 values)", ""),
                    ("nationality", "VARCHAR(100)", ""),
                    ("face_shape", "ENUM(7 shapes)", ""),
                    ("skin_tone", "ENUM(6 tones)", ""), ("skin_undertone", "ENUM(3)", ""),
                    ("personal_color", "ENUM(4 seasons)", ""), ("palette_id", "INT", "FK"),
                    ("confidence_score", "DECIMAL(5,2)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "analysis_reviews": {
                "class": "AnalysisReview", "file": "models/analysis.py", "pos": (16, 14),
                "cols": [
                    ("review_id", "INT", "PK,AI"), ("analysis_id", "INT", "FK,NN"),
                    ("user_id", "INT", "FK,NN"), ("is_accurate", "INT", ""),
                    ("comment", "TEXT", ""),
                ],
            },
        },
        "rels": [
            ("color_palettes", "analysis_results", "1", "*", "right", "left"),
            ("analysis_results", "analysis_reviews", "1", "*", "right", "left"),
        ],
    },
    {
        "title": "กลุ่มที่ 4: Recommendation Engine",
        "tables": {
            "users": {
                "class": "User", "file": "ref", "pos": (1, 8),
                "cols": [("user_id", "INT", "PK")],
            },
            "products": {
                "class": "Product", "file": "ref", "pos": (8, 8),
                "cols": [("product_id", "INT", "PK")],
            },
            "recommendation_rules": {
                "class": "RecommendationRule", "file": "models/recommendation.py", "pos": (1, 13),
                "cols": [
                    ("rule_id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("face_shape", "ENUM(7+any)", ""), ("skin_tone", "ENUM(6+any)", ""),
                    ("personal_color", "ENUM(4+any)", ""), ("priority", "INT", ""),
                ],
            },
            "recommendations": {
                "class": "Recommendation", "file": "models/recommendation.py", "pos": (8, 13),
                "cols": [
                    ("recommendation_id", "INT", "PK,AI"), ("analysis_id", "INT", "FK,NN"),
                    ("product_id", "INT", "FK,NN"), ("score", "DECIMAL(5,2)", ""),
                    ("created_at", "DATETIME", ""),
                ],
            },
            "recommendation_feedback": {
                "class": "RecommendationFeedback", "file": "models/recommendation.py", "pos": (15, 13),
                "cols": [
                    ("feedback_id", "INT", "PK,AI"), ("recommendation_id", "INT", "FK,NN"),
                    ("user_id", "INT", "FK,NN"), ("rating", "ENUM('like','dislike')", "NN"),
                ],
            },
            "favorites": {
                "class": "Favorite", "file": "models/recommendation.py", "pos": (15, 8),
                "cols": [
                    ("favorite_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("product_id", "INT", "FK,NN"), ("created_at", "DATETIME", ""),
                ],
            },
        },
        "rels": [
            ("recommendation_rules", "recommendations", "1..*", "1", "right", "left"),
            ("recommendations", "recommendation_feedback", "1", "*", "right", "left"),
            ("users", "recommendation_feedback", "1", "*", "right", "left"),
            ("users", "favorites", "1", "*", "right", "left"),
            ("products", "recommendation_rules", "1", "*", "bottom", "top"),
            ("products", "recommendations", "1", "*", "bottom", "top"),
            ("products", "favorites", "1", "*", "right", "left"),
        ],
    },
    {
        "title": "กลุ่มที่ 5: Skin Concerns",
        "tables": {
            "users": {
                "class": "User", "file": "ref", "pos": (1, 8),
                "cols": [("user_id", "INT", "PK")],
            },
            "products": {
                "class": "Product", "file": "ref", "pos": (15, 8),
                "cols": [("product_id", "INT", "PK")],
            },
            "skin_concerns": {
                "class": "SkinConcern", "file": "models/misc.py", "pos": (8, 12),
                "cols": [
                    ("concern_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""), ("icon_url", "VARCHAR(255)", ""),
                ],
            },
            "user_skin_concerns": {
                "class": "UserSkinConcern", "file": "models/misc.py", "pos": (1, 12),
                "cols": [
                    ("id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("concern_id", "INT", "FK,NN"), ("severity", "ENUM(3 levels)", ""),
                ],
            },
            "product_concerns": {
                "class": "ProductConcern", "file": "models/misc.py", "pos": (15, 12),
                "cols": [
                    ("id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("concern_id", "INT", "FK,NN"),
                ],
            },
        },
        "rels": [
            ("skin_concerns", "user_skin_concerns", "1", "*", "left", "right"),
            ("skin_concerns", "product_concerns", "1", "*", "right", "left"),
            ("users", "user_skin_concerns", "1", "*", "bottom", "top"),
            ("products", "product_concerns", "1", "*", "bottom", "top"),
        ],
    },
    {
        "title": "กลุ่มที่ 6: Gemini AI Chat",
        "tables": {
            "gemini_sessions": {
                "class": "GeminiSession", "file": "models/gemini.py", "pos": (1, 12),
                "cols": [
                    ("session_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("analysis_id", "INT", "FK"), ("title", "VARCHAR(200)", ""),
                ],
            },
            "gemini_messages": {
                "class": "GeminiMessage", "file": "models/gemini.py", "pos": (9, 12),
                "cols": [
                    ("message_id", "INT", "PK,AI"), ("session_id", "INT", "FK,NN"),
                    ("role", "ENUM('user','model')", "NN"), ("prompt", "TEXT", ""),
                    ("response", "TEXT", ""), ("image_input", "VARCHAR(255)", ""),
                    ("image_output", "VARCHAR(255)", ""),
                ],
            },
        },
        "rels": [("gemini_sessions", "gemini_messages", "1", "*", "right", "left")],
    },
    {
        "title": "กลุ่มที่ 7: Photo Editor & Saved Looks",
        "tables": {
            "users": {
                "class": "User", "file": "ref", "pos": (4.5, 17),
                "cols": [("user_id", "INT", "PK")],
            },
            "photo_edits": {
                "class": "PhotoEdit", "file": "models/photo_editor.py", "pos": (1, 13),
                "cols": [
                    ("edit_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("original_image", "VARCHAR(255)", ""), ("edited_image", "VARCHAR(255)", ""),
                    ("edit_config", "JSON", ""), ("source", "ENUM(3 values)", ""),
                ],
            },
            "saved_looks": {
                "class": "SavedLook", "file": "models/saved_look.py", "pos": (8, 13),
                "cols": [
                    ("look_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("name", "VARCHAR(100)", "NN"), ("category", "VARCHAR(50)", ""),
                    ("makeup_data", "JSON", "NN"), ("filter_data", "JSON", ""),
                ],
            },
            "edit_filters": {
                "class": "EditFilter", "file": "models/photo_editor.py", "pos": (15, 13),
                "cols": [
                    ("filter_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("category", "ENUM(5 types)", "NN"), ("config", "JSON", ""),
                    ("is_active", "INT", ""),
                ],
            },
            "edit_stickers": {
                "class": "EditSticker", "file": "models/photo_editor.py", "pos": (15, 8),
                "cols": [
                    ("sticker_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("category", "ENUM(4 types)", "NN"), ("image_url", "VARCHAR(255)", ""),
                ],
            },
        },
        "rels": [
            ("users", "photo_edits", "1", "*", "bottom", "top"),
            ("users", "saved_looks", "1", "*", "bottom", "top"),
        ],
    },
    {
        "title": "กลุ่มที่ 8: Blog, Reviews & Affiliate",
        "tables": {
            "users": {
                "class": "User", "file": "ref", "pos": (8, 18),
                "cols": [("user_id", "INT", "PK")],
            },
            "products": {
                "class": "Product", "file": "ref", "pos": (15, 18),
                "cols": [("product_id", "INT", "PK")],
            },
            "product_links": {
                "class": "ProductLink", "file": "ref", "pos": (1, 5),
                "cols": [("link_id", "INT", "PK")],
            },
            "blog_categories": {
                "class": "BlogCategory", "file": "models/blog.py", "pos": (1, 14),
                "cols": [
                    ("category_id", "INT", "PK,AI"), ("name", "VARCHAR(100)", "NN"),
                    ("description", "TEXT", ""),
                ],
            },
            "blog_posts": {
                "class": "BlogPost", "file": "models/blog.py", "pos": (8, 14),
                "cols": [
                    ("post_id", "INT", "PK,AI"), ("category_id", "INT", "FK"),
                    ("author_id", "INT", "FK,NN"), ("title", "VARCHAR(200)", "NN"),
                    ("slug", "VARCHAR(200)", "UK"), ("content", "TEXT", ""),
                    ("views", "INT", ""), ("is_published", "INT", ""),
                ],
            },
            "product_reviews": {
                "class": "ProductReview", "file": "models/misc.py", "pos": (15, 14),
                "cols": [
                    ("review_id", "INT", "PK,AI"), ("product_id", "INT", "FK,NN"),
                    ("user_id", "INT", "FK,NN"), ("rating", "INT", "NN"),
                    ("comment", "TEXT", ""), ("platform", "ENUM(4)", ""),
                ],
            },
            "click_logs": {
                "class": "ClickLog", "file": "models/commission.py", "pos": (1, 8.5),
                "cols": [
                    ("log_id", "INT", "PK,AI"), ("link_id", "INT", "FK,NN"),
                    ("user_id", "INT", "FK"), ("platform", "ENUM(3)", "NN"),
                    ("ip_address", "VARCHAR(45)", ""),
                ],
            },
            "commissions": {
                "class": "Commission", "file": "models/commission.py", "pos": (8, 8.5),
                "cols": [
                    ("commission_id", "INT", "PK,AI"), ("link_id", "INT", "FK,NN"),
                    ("user_id", "INT", "FK"), ("amount", "DECIMAL(10,2)", ""),
                    ("status", "ENUM(3 statuses)", ""),
                ],
            },
        },
        "rels": [
            ("blog_categories", "blog_posts", "1", "*", "right", "left"),
            ("users", "blog_posts", "1", "*", "bottom", "top"),
            ("users", "product_reviews", "1", "*", "bottom", "top"),
            ("products", "product_reviews", "1", "*", "bottom", "top"),
            ("product_links", "click_logs", "1", "*", "bottom", "top"),
            ("product_links", "commissions", "1", "*", "bottom", "top"),
        ],
    },
    {
        "title": "กลุ่มที่ 9: Admin, Notifications & Analytics",
        "tables": {
            "users": {
                "class": "User", "file": "ref", "pos": (8, 17),
                "cols": [("user_id", "INT", "PK")],
            },
            "notifications": {
                "class": "Notification", "file": "models/misc.py", "pos": (1, 13),
                "cols": [
                    ("notification_id", "INT", "PK,AI"), ("user_id", "INT", "FK,NN"),
                    ("title", "VARCHAR(200)", "NN"), ("message", "TEXT", ""),
                    ("type", "ENUM(4 types)", ""), ("is_read", "INT", ""),
                ],
            },
            "admin_logs": {
                "class": "AdminLog", "file": "models/misc.py", "pos": (8, 13),
                "cols": [
                    ("log_id", "INT", "PK,AI"), ("admin_id", "INT", "FK,NN"),
                    ("action", "VARCHAR(100)", "NN"), ("table_name", "VARCHAR(100)", ""),
                    ("old_value", "JSON", ""), ("new_value", "JSON", ""),
                ],
            },
            "user_behaviors": {
                "class": "UserBehavior", "file": "models/behavior.py", "pos": (15, 13),
                "cols": [
                    ("id", "INT", "PK,AI"), ("user_id", "INT", "FK"),
                    ("session_id", "VARCHAR(100)", ""), ("event_type", "VARCHAR(50)", "NN"),
                    ("event_data", "JSON", ""), ("page", "VARCHAR(50)", ""),
                ],
            },
            "search_history": {
                "class": "SearchHistory", "file": "models/misc.py", "pos": (1, 7.5),
                "cols": [
                    ("search_id", "INT", "PK,AI"), ("user_id", "INT", "FK"),
                    ("keyword", "VARCHAR(255)", "NN"),
                ],
            },
            "banners": {
                "class": "Banner", "file": "models/misc.py", "pos": (15, 7.5),
                "cols": [
                    ("banner_id", "INT", "PK,AI"), ("title", "VARCHAR(200)", ""),
                    ("image_url", "VARCHAR(255)", "NN"), ("position", "ENUM(3)", ""),
                    ("is_active", "INT", ""),
                ],
            },
        },
        "rels": [
            ("users", "notifications", "1", "*", "bottom", "top"),
            ("users", "admin_logs", "1", "*", "bottom", "top"),
            ("users", "user_behaviors", "1", "*", "bottom", "top"),
            ("users", "search_history", "1", "*", "bottom", "top"),
        ],
    },
    {
        "title": "กลุ่มที่ 10: ภาพรวมความสัมพันธ์ Core Entities",
        "tables": {
            "users": {
                "class": "User", "file": "models/user.py", "pos": (1, 16),
                "cols": [("user_id", "INT", "PK"), ("username", "VARCHAR(50)", "UK")],
            },
            "analysis_results": {
                "class": "AnalysisResult", "file": "models/analysis.py", "pos": (8, 16),
                "cols": [("analysis_id", "INT", "PK"), ("user_id", "INT", "FK"),
                         ("face_shape", "ENUM", ""), ("personal_color", "ENUM", "")],
            },
            "color_palettes": {
                "class": "ColorPalette", "file": "models/analysis.py", "pos": (15, 16),
                "cols": [("palette_id", "INT", "PK"), ("season", "ENUM", ""),
                         ("best_colors", "JSON", "")],
            },
            "recommendations": {
                "class": "Recommendation", "file": "models/recommendation.py", "pos": (8, 10.5),
                "cols": [("recommendation_id", "INT", "PK"), ("score", "DECIMAL", "")],
            },
            "products": {
                "class": "Product", "file": "models/product.py", "pos": (15, 10.5),
                "cols": [("product_id", "INT", "PK"), ("name", "VARCHAR(200)", ""),
                         ("price", "DECIMAL", "")],
            },
            "gemini_sessions": {
                "class": "GeminiSession", "file": "models/gemini.py", "pos": (1, 10.5),
                "cols": [("session_id", "INT", "PK"), ("title", "VARCHAR(200)", "")],
            },
            "favorites": {
                "class": "Favorite", "file": "models/recommendation.py", "pos": (8, 5.5),
                "cols": [("favorite_id", "INT", "PK")],
            },
            "photo_edits": {
                "class": "PhotoEdit", "file": "models/photo_editor.py", "pos": (1, 5.5),
                "cols": [("edit_id", "INT", "PK"), ("source", "ENUM", "")],
            },
        },
        "rels": [
            ("users", "analysis_results", "1", "*", "right", "left"),
            ("analysis_results", "color_palettes", "*", "1", "right", "left"),
            ("analysis_results", "recommendations", "1", "*", "bottom", "top"),
            ("products", "recommendations", "1", "*", "left", "right"),
            ("users", "gemini_sessions", "1", "*", "bottom", "top"),
            ("users", "favorites", "1", "*", "bottom", "top"),
            ("products", "favorites", "1", "*", "bottom", "top"),
            ("users", "photo_edits", "1", "*", "bottom", "top"),
        ],
    },
]


# ═══════════════════════════════════════
# Drawing engine (same as ER diagram)
# ═══════════════════════════════════════

def _box_h(cols):
    return HEADER_H + max(len(cols), 1) * LINE_H + 0.35

def _draw_table(ax, tbl_name, data):
    x, y_top = data["pos"]
    cols = data["cols"]
    cls_name = data["class"]
    w, h = TBL_W, _box_h(cols)
    rect = FancyBboxPatch((x, y_top - h), w, h, boxstyle="round,pad=0.05",
                           facecolor=_hex(TBL_FILL), edgecolor=_hex(TBL_BORDER),
                           linewidth=2.0, zorder=2)
    ax.add_patch(rect)
    ax.text(x + w/2, y_top - 0.22, tbl_name, ha="center", va="center",
            fontsize=TBL_NAME_FS, fontweight="bold", fontfamily=FONT,
            color=_hex(TBL_BORDER), zorder=4)
    ax.text(x + w/2, y_top - 0.52, f"({cls_name})", ha="center", va="center",
            fontsize=COL_FS, fontstyle="italic", fontfamily=FONT,
            color=_hex("#616161"), zorder=4)
    sep_y = y_top - HEADER_H
    ax.plot([x+0.08, x+w-0.08], [sep_y, sep_y], color=_hex(TBL_BORDER), linewidth=1.2, zorder=3)
    cy = sep_y - 0.25
    for col_name, col_type, flags in cols:
        fs = set(f.strip() for f in flags.split(",")) if flags else set()
        if "PK" in fs:
            icon, nc = "PK ", _hex(PK_COLOR)
        elif "FK" in fs:
            icon, nc = "FK ", _hex(FK_COLOR)
        else:
            icon, nc = "     ", _hex(COL_COLOR)
        badge = ""
        for tag in ["PK","FK","UK","NN","AI"]:
            if tag in fs: badge += f" [{tag}]"
        ax.text(x + PAD_L, cy, f"{icon}{col_name} : {col_type}{badge}",
                ha="left", va="center", fontsize=COL_FS, fontfamily=FONT, color=nc, zorder=4)
        cy -= LINE_H
    data["_bbox"] = (x, y_top - h, x + w, y_top)

def _anchor(bbox, side):
    x0, y0, x1, y1 = bbox
    cx, cy = (x0+x1)/2, (y0+y1)/2
    return {"top":(cx,y1),"bottom":(cx,y0),"left":(x0,cy),"right":(x1,cy)}[side]

def _draw_crow_foot(ax, pos, side):
    x, y = pos; sz = 0.18
    pts = {
        "left": ([x-sz,x,x-sz],[y+sz*.7,y,y-sz*.7]),
        "right": ([x+sz,x,x+sz],[y+sz*.7,y,y-sz*.7]),
        "top": ([x-sz*.7,x,x+sz*.7],[y+sz,y,y+sz]),
        "bottom": ([x-sz*.7,x,x+sz*.7],[y-sz,y,y-sz]),
    }
    ax.plot(*pts[side], color=_hex(REL_COLOR), linewidth=1.5, zorder=6)

def _draw_rel(ax, td, ta, tb, ma, mb, sa, sb):
    ba, bb = td[ta]["_bbox"], td[tb]["_bbox"]
    pa, pb = list(_anchor(ba, sa)), list(_anchor(bb, sb))
    if sa in ("top","bottom") and sb in ("top","bottom"):
        my = (pa[1]+pb[1])/2
        ax.plot([pa[0],pa[0],pb[0],pb[0]],[pa[1],my,my,pb[1]], color=_hex(REL_COLOR), lw=1.5, zorder=1)
    elif sa in ("left","right") and sb in ("left","right"):
        mx = (pa[0]+pb[0])/2
        ax.plot([pa[0],mx,mx,pb[0]],[pa[1],pa[1],pb[1],pb[1]], color=_hex(REL_COLOR), lw=1.5, zorder=1)
    else:
        if sa in ("top","bottom"):
            ax.plot([pa[0],pa[0],pb[0]],[pa[1],pb[1],pb[1]], color=_hex(REL_COLOR), lw=1.5, zorder=1)
        else:
            ax.plot([pa[0],pb[0],pb[0]],[pa[1],pa[1],pb[1]], color=_hex(REL_COLOR), lw=1.5, zorder=1)
    if mb == "*": _draw_crow_foot(ax, pb, sb)
    d = 0.35
    def _mp(p, s):
        return {"top":(p[0]+.2,p[1]+d),"bottom":(p[0]+.2,p[1]-d),
                "right":(p[0]+d,p[1]+.22),"left":(p[0]-d-.1,p[1]+.22)}[s]
    if ma: ax.text(*_mp(pa,sa), ma, fontsize=11, fontweight="bold", color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)
    if mb: ax.text(*_mp(pb,sb), mb, fontsize=11, fontweight="bold", color=_hex(MULT_COLOR), ha="center", va="center", zorder=5)

def render_group(group):
    tbls = group["tables"]; rels = group["rels"]
    ax_vals = [d["pos"] for d in tbls.values()]
    all_x = [p[0] for p in ax_vals]; all_y = [p[1] for p in ax_vals]
    mn_x, mx_x = min(all_x)-1.5, max(all_x)+TBL_W+2
    mn_y = min(all_y) - max(_box_h(d["cols"]) for d in tbls.values()) - 1.5
    mx_y = max(all_y) + 2
    fw, fh = max(mx_x-mn_x, 12), max(mx_y-mn_y, 7)
    fig, ax = plt.subplots(figsize=(fw, fh))
    ax.set_xlim(mn_x, mx_x); ax.set_ylim(mn_y, mx_y)
    ax.set_aspect("equal"); ax.axis("off"); fig.patch.set_facecolor("white")
    for n, d in tbls.items(): _draw_table(ax, n, d)
    for r in rels: _draw_rel(ax, tbls, *r)
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig); buf.seek(0)
    return buf


# ═══════════════════════════════════════
# Helper: add styled paragraph
# ═══════════════════════════════════════

def add_heading_th(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    elif level == 2:
        r = p.add_run(text)
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    elif level == 3:
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def add_body_no_indent(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


# ═══════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════

doc = Document()

# Portrait A4
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# ════════════════════════════════════════════════
# SECTION 1: บทนำ
# ════════════════════════════════════════════════

add_heading_th(doc, "การแปลงคลาสเป็นตารางในฐานข้อมูลเชิงสัมพันธ์", 1)

add_body(doc,
    "ในการพัฒนาระบบซอฟต์แวร์แบบ Object-Oriented ที่ใช้ฐานข้อมูลเชิงสัมพันธ์ (Relational Database) เป็นตัวจัดเก็บข้อมูล "
    "จำเป็นต้องมีกระบวนการ \"แปลง\" (Mapping) จากโมเดลเชิงวัตถุ (Class) ไปเป็นโครงสร้างตาราง (Table) ในฐานข้อมูล "
    "กระบวนการนี้เรียกว่า Object-Relational Mapping (ORM)")

add_body(doc,
    "โปรเจกต์ AuraMatch ใช้ SQLAlchemy ซึ่งเป็น ORM Framework สำหรับ Python ในการนิยามคลาส (Declarative Base) "
    "แล้วแปลงเป็นตารางในฐานข้อมูล MySQL โดยอัตโนมัติ โดยมีคลาสทั้งหมด 36 คลาส แปลงเป็น 36 ตาราง "
    "จัดกลุ่มเป็น 9 โดเมน")


# ════════════════════════════════════════════════
# SECTION 2: หลักการแปลง
# ════════════════════════════════════════════════

add_heading_th(doc, "1. หลักการแปลงคลาสเป็นตารางในฐานข้อมูลเชิงสัมพันธ์", 1)

add_heading_th(doc, "1.1 แปลงคลาส (Class) เป็นตาราง (Table)", 2)

add_body(doc,
    "คลาสแต่ละตัวในโมเดลเชิงวัตถุจะถูกแปลงเป็นตาราง (Table) หนึ่งตารางในฐานข้อมูล "
    "โดยชื่อตารางจะตั้งตามแบบแผน snake_case เช่น คลาส UserProfile จะถูกแปลงเป็นตาราง user_profiles "
    "และคลาส AnalysisResult จะถูกแปลงเป็นตาราง analysis_results")

add_body_no_indent(doc, "ตัวอย่างจากโปรเจกต์ AuraMatch:")

# Table: class → table example
t1 = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
for i, h in enumerate(["ORM Class (Python)", "Table Name (MySQL)", "หมายเหตุ"]):
    t1.rows[0].cells[i].text = h
    for p in t1.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
examples_1 = [
    ("User", "users", "ศูนย์กลางของระบบ — มี FK จาก 16 ตารางอื่น"),
    ("UserProfile", "user_profiles", "ข้อมูลโปรไฟล์ผู้ใช้ (1:1 กับ users)"),
    ("Product", "products", "สินค้าเครื่องสำอาง"),
    ("AnalysisResult", "analysis_results", "ผลวิเคราะห์ใบหน้าและ Personal Color"),
    ("ColorPalette", "color_palettes", "พาเลตต์สี 12 sub-seasons (JSON)"),
    ("GeminiSession", "gemini_sessions", "เซสชันแชท AI"),
    ("SavedLook", "saved_looks", "ลุคแต่งหน้าที่บันทึก"),
]
for cls, tbl, note in examples_1:
    row = t1.add_row().cells
    row[0].text = cls; row[1].text = tbl; row[2].text = note
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)

doc.add_paragraph("")

# ── 1.2 Attribute → Column ──
add_heading_th(doc, "1.2 แปลง Attribute เป็น Column", 2)

add_body(doc,
    "แอตทริบิวต์ (Attribute) ของคลาสแต่ละตัวจะถูกแปลงเป็นคอลัมน์ (Column) ในตาราง "
    "โดยชนิดข้อมูลเชิงวัตถุจะถูกแปลงเป็น Data Type ของฐานข้อมูล ดังนี้:")

t2 = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
for i, h in enumerate(["Python Type (SQLAlchemy)", "MySQL Data Type", "ตัวอย่าง Attribute", "ตัวอย่าง Column"]):
    t2.rows[0].cells[i].text = h
    for p in t2.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
type_maps = [
    ("Integer", "INT", "user_id", "user_id INT"),
    ("String(n)", "VARCHAR(n)", "username: String(50)", "username VARCHAR(50)"),
    ("Text", "TEXT", "description", "description TEXT"),
    ("DECIMAL(p,s)", "DECIMAL(p,s)", "price: DECIMAL(10,2)", "price DECIMAL(10,2)"),
    ("DateTime", "DATETIME", "created_at", "created_at DATETIME"),
    ("Date", "DATE", "birth_date", "birth_date DATE"),
    ("Enum(...)", "ENUM(...)", "role: Enum('admin','user')", "role ENUM('admin','user')"),
    ("JSON", "JSON", "best_colors", "best_colors JSON"),
]
for py, sql, attr, col in type_maps:
    row = t2.add_row().cells
    row[0].text = py; row[1].text = sql; row[2].text = attr; row[3].text = col
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)

doc.add_paragraph("")

# ── 1.3 Primary Key ──
add_heading_th(doc, "1.3 แปลง Object Identity เป็น Primary Key", 2)

add_body(doc,
    "ในระบบเชิงวัตถุ แต่ละอ็อบเจกต์มี Identity เฉพาะตัว เมื่อแปลงเป็นตารางจะใช้ Primary Key (PK) เป็นตัวระบุแถว "
    "โปรเจกต์ AuraMatch ใช้ Surrogate Key แบบ Auto Increment ทุกตาราง เช่น user_id, product_id, analysis_id "
    "เพราะไม่ต้องพึ่งพาข้อมูลทางธุรกิจ (Natural Key) ที่อาจเปลี่ยนแปลง")

add_body_no_indent(doc, "นอกจากนี้ยังกำหนด Unique Key (UK) เพื่อรับรองความไม่ซ้ำเชิงธุรกิจ เช่น:")
add_bullet(doc, "users.username — ชื่อผู้ใช้ห้ามซ้ำ")
add_bullet(doc, "users.email — อีเมลห้ามซ้ำ")
add_bullet(doc, "product_tags.name — ชื่อแท็กห้ามซ้ำ")
add_bullet(doc, "favorites (user_id, product_id) — Compound Unique Key ป้องกันกดซ้ำ")
add_bullet(doc, "user_profiles.user_id — Unique เพื่อบังคับ 1:1")

doc.add_paragraph("")

# ── 1.4 Relationships ──
add_heading_th(doc, "1.4 แปลงความสัมพันธ์ (Relationship) เป็น Foreign Key", 2)

add_body(doc,
    "ความสัมพันธ์ระหว่างคลาสจะถูกแปลงเป็น Foreign Key (FK) ในตารางฐานข้อมูล "
    "โดยมี 3 รูปแบบหลัก ดังนี้:")

add_heading_th(doc, "1.4.1 One-to-One (1:1)", 3)
add_body(doc,
    "คลาสที่มีความสัมพันธ์แบบหนึ่งต่อหนึ่ง จะแปลงโดยเพิ่ม FK พร้อม Unique Constraint ในตารางฝั่งหนึ่ง "
    "เพื่อบังคับว่าแต่ละแถวสัมพันธ์กับอีกตารางได้เพียง 1 แถว")
add_body_no_indent(doc, "ตัวอย่าง: User ↔ UserProfile")
add_bullet(doc, "ตาราง user_profiles มีคอลัมน์ user_id เป็น FK อ้างอิงไปยัง users.user_id")
add_bullet(doc, "กำหนด Unique Constraint บน user_profiles.user_id เพื่อบังคับ 1:1")
add_bullet(doc, "ON DELETE CASCADE — เมื่อลบ User จะลบ Profile ตามไปด้วย")

add_heading_th(doc, "1.4.2 One-to-Many (1:N)", 3)
add_body(doc,
    "เป็นรูปแบบที่พบมากที่สุดในโปรเจกต์ (35 ความสัมพันธ์) โดยเพิ่ม FK ในตารางฝั่ง \"Many\" "
    "อ้างอิงไปยัง PK ของตารางฝั่ง \"One\"")
add_body_no_indent(doc, "ตัวอย่าง:")
add_bullet(doc, "User (1) → AnalysisResult (N): ผู้ใช้ 1 คน วิเคราะห์ได้หลายครั้ง → analysis_results.user_id FK")
add_bullet(doc, "Product (1) → ProductLink (N): สินค้า 1 ชิ้น มีลิงก์ขายหลายแพลตฟอร์ม → product_links.product_id FK")
add_bullet(doc, "GeminiSession (1) → GeminiMessage (N): เซสชัน 1 ครั้ง มีหลายข้อความ → gemini_messages.session_id FK")

add_heading_th(doc, "1.4.3 Many-to-Many (M:N)", 3)
add_body(doc,
    "ความสัมพันธ์แบบกลุ่มต่อกลุ่มไม่สามารถแทนด้วย FK เดียวได้ ต้องสร้าง \"ตารางเชื่อม\" (Junction Table / Associative Table) "
    "ที่มี FK 2 ตัวชี้ไปยังทั้งสองฝั่ง")
add_body_no_indent(doc, "ตัวอย่างจากโปรเจกต์:")

t_mn = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
for i, h in enumerate(["ตาราง A", "ตาราง B", "Junction Table", "หมายเหตุ"]):
    t_mn.rows[0].cells[i].text = h
    for p in t_mn.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
mn_data = [
    ("products", "product_tags", "product_tag_map", "สินค้า ↔ แท็ก"),
    ("users", "skin_concerns", "user_skin_concerns", "ผู้ใช้ ↔ ปัญหาผิว (มี severity เป็น attribute เพิ่ม)"),
    ("products", "skin_concerns", "product_concerns", "สินค้า ↔ ปัญหาผิวที่รักษาได้"),
]
for a, b, jt, note in mn_data:
    row = t_mn.add_row().cells
    row[0].text = a; row[1].text = b; row[2].text = jt; row[3].text = note
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)

doc.add_paragraph("")

# ── 1.5 ON DELETE ──
add_heading_th(doc, "1.5 กฎการลบ (Referential Integrity — ON DELETE)", 2)

add_body(doc,
    "เมื่อแถวในตารางแม่ถูกลบ จะส่งผลต่อแถวในตารางลูกที่อ้างอิง FK อยู่ "
    "โปรเจกต์ AuraMatch กำหนดนโยบายการลบ 2 แบบ ดังนี้:")

t_del = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
for i, h in enumerate(["ON DELETE", "ความหมาย", "ตัวอย่างในโปรเจกต์"]):
    t_del.rows[0].cells[i].text = h
    for p in t_del.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
del_data = [
    ("CASCADE", "ลบแถวลูกตามไปด้วย", "ลบ User → ลบ AnalysisResult, Favorite, PhotoEdit ตาม"),
    ("SET NULL", "เปลี่ยน FK เป็น NULL", "ลบ Brand → products.brand_id เปลี่ยนเป็น NULL (สินค้ายังอยู่)"),
]
for od, mean, ex in del_data:
    row = t_del.add_row().cells
    row[0].text = od; row[1].text = mean; row[2].text = ex
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)

doc.add_paragraph("")

# ── 1.6 Enum ──
add_heading_th(doc, "1.6 แปลง Enumeration เป็น ENUM Type", 2)

add_body(doc,
    "คลาสที่มีแอตทริบิวต์แบบ Enumeration (ค่าจำกัดจำนวนตัวเลือก) จะถูกแปลงเป็น ENUM Data Type ของ MySQL "
    "ซึ่งจำกัดค่าที่บันทึกได้ในระดับฐานข้อมูล โปรเจกต์ AuraMatch มี ENUM จำนวนมาก เช่น:")

t_enum = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
for i, h in enumerate(["ตาราง.คอลัมน์", "ค่าที่เป็นไปได้", "ความหมาย"]):
    t_enum.rows[0].cells[i].text = h
    for p in t_enum.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
enum_data = [
    ("users.role", "admin, user", "บทบาทผู้ใช้"),
    ("analysis_results.face_shape", "oval, round, square, heart,\noblong, diamond, triangle", "รูปทรงใบหน้า 7 แบบ"),
    ("analysis_results.skin_tone", "fair, light, medium, tan, dark, deep", "โทนสีผิว 6 ระดับ"),
    ("analysis_results.personal_color", "spring, summer, autumn, winter", "Personal Color 4 ฤดู"),
    ("product_links.platform", "shopee, tiktok, lazada,\nsephora, watsons, ulta", "แพลตฟอร์มที่ขาย"),
    ("commissions.status", "pending, confirmed, paid", "สถานะค่าคอมมิชชั่น"),
    ("recommendation_feedback.rating", "like, dislike", "ความพึงพอใจ"),
]
for col, vals, mean in enum_data:
    row = t_enum.add_row().cells
    row[0].text = col; row[1].text = vals; row[2].text = mean
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)

doc.add_paragraph("")

# ── 1.7 JSON ──
add_heading_th(doc, "1.7 แปลง Complex Object เป็น JSON Column", 2)

add_body(doc,
    "แอตทริบิวต์ที่เป็น Object ซ้อน (Nested Object) หรือ Collection ที่ไม่ต้องการค้นหาแยก "
    "จะถูกแปลงเป็นคอลัมน์ชนิด JSON แทนการสร้างตารางใหม่ เพื่อลดจำนวน JOIN ในการ Query:")

t_json = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
for i, h in enumerate(["ตาราง.คอลัมน์", "โครงสร้าง JSON", "เหตุผลที่ใช้ JSON"]):
    t_json.rows[0].cells[i].text = h
    for p in t_json.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(10)
json_data = [
    ("color_palettes.best_colors", '[{"L": 75, "a": 12, "b": 8}, ...]', "เก็บค่าสี CIELAB หลายสีต่อ 1 palette"),
    ("color_palettes.avoid_colors", '[{"L": 30, "a": -5, "b": 2}, ...]', "สีที่ควรหลีกเลี่ยงของแต่ละ palette"),
    ("saved_looks.makeup_data", '{"lipstick": {...}, "blush": {...}}', "สถานะแต่งหน้าทั้งหมดเป็น 1 ก้อน"),
    ("saved_looks.filter_data", '{"brightness": 1.2, "contrast": 1.0}', "ค่าฟิลเตอร์เป็นชุดเดียว"),
    ("admin_logs.old_value", '{"price": 590}', "Audit log — เก็บค่าเก่าก่อนแก้ไข"),
    ("admin_logs.new_value", '{"price": 490}', "Audit log — เก็บค่าใหม่หลังแก้ไข"),
    ("user_behaviors.event_data", '{"product_id": 42, "query": "..."}', "ข้อมูลเสริมของ Event แต่ละประเภท"),
]
for col, structure, reason in json_data:
    row = t_json.add_row().cells
    row[0].text = col; row[1].text = structure; row[2].text = reason
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(10)


# ════════════════════════════════════════════════
# SECTION 3: ตารางสรุป
# ════════════════════════════════════════════════

doc.add_page_break()
add_heading_th(doc, "2. ตารางสรุปการแปลง ORM Class เป็น Relational Table ทั้ง 36 ตาราง", 1)

add_body(doc,
    "ตารางด้านล่างแสดงการ Mapping คลาสทั้ง 36 คลาสในโปรเจกต์ AuraMatch "
    "ไปยังตารางในฐานข้อมูล MySQL พร้อมระบุไฟล์ต้นทาง จำนวนคอลัมน์ และกลุ่มโดเมน:")

doc.add_paragraph("")

t_all = doc.add_table(rows=1, cols=6, style="Light Grid Accent 1")
for i, h in enumerate(["#", "ORM Class", "Table Name", "Source File", "Columns", "Domain"]):
    t_all.rows[0].cells[i].text = h
    for p in t_all.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(9)

idx = 0
for gi, g in enumerate(ER_GROUPS):
    if gi == len(ER_GROUPS) - 1: continue  # skip overview
    domain = g["title"].split(":")[1].strip() if ":" in g["title"] else g["title"]
    for tbl_name, tbl_data in g["tables"].items():
        if tbl_data.get("file") == "ref": continue  # skip reference entities
        idx += 1
        row = t_all.add_row().cells
        row[0].text = str(idx)
        row[1].text = tbl_data["class"]
        row[2].text = tbl_name
        row[3].text = tbl_data["file"]
        row[4].text = str(len(tbl_data["cols"]))
        row[5].text = domain.split("(")[0].strip() if "(" in domain else domain
        for c in row:
            for p in c.paragraphs:
                for rr in p.runs: rr.font.size = Pt(9)


# ════════════════════════════════════════════════
# SECTION 4: ตาราง FK ทั้งหมด
# ════════════════════════════════════════════════

doc.add_page_break()
add_heading_th(doc, "3. ตารางสรุป Foreign Key ทั้งหมด (40 ความสัมพันธ์)", 1)

add_body(doc,
    "ตารางด้านล่างแสดง Foreign Key ทั้งหมดในฐานข้อมูล พร้อมระบุตารางแม่ (Parent), ตารางลูก (Child), "
    "คอลัมน์ FK, ประเภทความสัมพันธ์ และนโยบาย ON DELETE:")

doc.add_paragraph("")

rtbl = doc.add_table(rows=1, cols=6, style="Light Grid Accent 1")
for i, h in enumerate(["#", "Parent (1)", "Child (N)", "FK Column", "Type", "ON DELETE"]):
    rtbl.rows[0].cells[i].text = h
    for p in rtbl.rows[0].cells[i].paragraphs:
        for rr in p.runs: rr.font.bold = True; rr.font.size = Pt(9)

all_fk = [
    ("users", "user_profiles", "user_id", "1:1", "CASCADE"),
    ("users", "password_resets", "user_id", "1:N", "CASCADE"),
    ("users", "analysis_results", "user_id", "1:N", "CASCADE"),
    ("users", "analysis_reviews", "user_id", "1:N", "CASCADE"),
    ("users", "favorites", "user_id", "1:N", "CASCADE"),
    ("users", "notifications", "user_id", "1:N", "CASCADE"),
    ("users", "gemini_sessions", "user_id", "1:N", "CASCADE"),
    ("users", "blog_posts", "author_id", "1:N", "-"),
    ("users", "photo_edits", "user_id", "1:N", "CASCADE"),
    ("users", "saved_looks", "user_id", "1:N", "CASCADE"),
    ("users", "product_reviews", "user_id", "1:N", "CASCADE"),
    ("users", "recommendation_feedback", "user_id", "1:N", "CASCADE"),
    ("users", "search_history", "user_id", "1:N", "SET NULL"),
    ("users", "user_skin_concerns", "user_id", "1:N", "CASCADE"),
    ("users", "admin_logs", "admin_id", "1:N", "-"),
    ("users", "user_behaviors", "user_id", "1:N", "SET NULL"),
    ("users", "click_logs", "user_id", "1:N", "SET NULL"),
    ("users", "commissions", "user_id", "1:N", "SET NULL"),
    ("brands", "products", "brand_id", "1:N", "SET NULL"),
    ("product_categories", "products", "category_id", "1:N", "SET NULL"),
    ("products", "product_links", "product_id", "1:N", "CASCADE"),
    ("products", "product_color_shades", "product_id", "1:N", "CASCADE"),
    ("products", "product_tag_map", "product_id", "1:N", "CASCADE"),
    ("product_tags", "product_tag_map", "tag_id", "1:N", "CASCADE"),
    ("products", "product_reviews", "product_id", "1:N", "CASCADE"),
    ("products", "favorites", "product_id", "1:N", "CASCADE"),
    ("products", "recommendation_rules", "product_id", "1:N", "CASCADE"),
    ("products", "recommendations", "product_id", "1:N", "-"),
    ("products", "product_concerns", "product_id", "1:N", "CASCADE"),
    ("color_palettes", "analysis_results", "palette_id", "1:N", "-"),
    ("analysis_results", "recommendations", "analysis_id", "1:N", "CASCADE"),
    ("analysis_results", "analysis_reviews", "analysis_id", "1:N", "CASCADE"),
    ("analysis_results", "gemini_sessions", "analysis_id", "1:N", "SET NULL"),
    ("recommendations", "recommendation_feedback", "recommendation_id", "1:N", "CASCADE"),
    ("product_links", "click_logs", "link_id", "1:N", "-"),
    ("product_links", "commissions", "link_id", "1:N", "-"),
    ("blog_categories", "blog_posts", "category_id", "1:N", "SET NULL"),
    ("skin_concerns", "user_skin_concerns", "concern_id", "1:N", "-"),
    ("skin_concerns", "product_concerns", "concern_id", "1:N", "-"),
    ("gemini_sessions", "gemini_messages", "session_id", "1:N", "CASCADE"),
]
for i, (parent, child, fk, rtype, ondel) in enumerate(all_fk, 1):
    row = rtbl.add_row().cells
    row[0].text = str(i); row[1].text = parent; row[2].text = child
    row[3].text = fk; row[4].text = rtype; row[5].text = ondel
    for c in row:
        for p in c.paragraphs:
            for rr in p.runs: rr.font.size = Pt(9)


# ════════════════════════════════════════════════
# SECTION 5: ER Diagram ประกอบ (landscape pages)
# ════════════════════════════════════════════════

doc.add_page_break()
add_heading_th(doc, "4. ER Diagram ประกอบ", 1)

add_body(doc,
    "ER Diagram (Entity-Relationship Diagram) ด้านล่างแสดงโครงสร้างตารางและความสัมพันธ์ "
    "ที่ได้จากการแปลง ORM Class เป็นตารางในฐานข้อมูลเชิงสัมพันธ์ แบ่งตามกลุ่มโดเมน "
    "โดยแต่ละกล่องตารางแสดงชื่อตาราง, ชื่อคลาสต้นทาง, คอลัมน์พร้อม Data Type และ Constraint (PK/FK/UK/NN)")

add_body_no_indent(doc, "สัญลักษณ์ในแผนภาพ:")
add_bullet(doc, "PK (สีน้ำเงิน) = Primary Key — คีย์หลัก")
add_bullet(doc, "FK (สีแดง) = Foreign Key — คีย์นอก อ้างอิงตารางอื่น")
add_bullet(doc, "UK = Unique Key — ค่าห้ามซ้ำ")
add_bullet(doc, "NN = Not Null — ห้ามเป็นค่าว่าง")
add_bullet(doc, "เส้นเชื่อม = ความสัมพันธ์ระหว่างตาราง")
add_bullet(doc, "ตัวเลข 1, * = Multiplicity (หนึ่ง, หลาย)")

# Switch to landscape for ER diagrams
for group in ER_GROUPS:
    # New landscape section
    new_section = doc.add_section(start_type=2)  # new page
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.top_margin = Cm(1.2)
    new_section.bottom_margin = Cm(1.2)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(group["title"])
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    img_buf = render_group(group)
    doc.add_picture(img_buf, width=Cm(25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_list = ", ".join(f"{n} ({d['class']})" for n, d in group["tables"].items())
    r = cap.add_run(f"Tables ({len(group['tables'])}): {tbl_list}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    r.font.italic = True


# ════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════

out_path = os.path.join(OUT_DIR, "AuraMatch_Class_to_Table_Report.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
