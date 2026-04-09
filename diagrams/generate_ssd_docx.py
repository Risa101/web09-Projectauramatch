"""
AuraMatch — System Sequence Diagrams (Word .docx)
One SSD per use case, landscape A4, readable text.
Actor ↔ System boundary only (black-box view).
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch
import io, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

FONT = "Tahoma"
OUT_DIR = "/Users/saridbutchuang/Desktop/web09_Projectauramatch/diagrams"

# ─── Colors ───
BG = "white"
ACTOR_HEAD = "#D6EEFF"
SYSTEM_HEAD = "#D7F5D7"
LIFELINE_C = "#9E9E9E"
ARROW_C = "#2C2C2C"
RETURN_C = "#1565C0"
MSG_FS = 10.5
LABEL_FS = 12
HEAD_FS = 13
NOTE_FS = 9.5
NOTE_BG = "#FFFDE7"
NOTE_BORDER = "#BDBDBD"

# ═══════════════════════════════════════════
# SSD Data: list of use cases
# Each message: (direction, label, note)
#   direction: "->": actor to system, "<-": system to actor, "note": note box
# ═══════════════════════════════════════════

SSDS = [
    {
        "title": "System Sequence Diagram: การลงทะเบียน (Register)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการลงทะเบียนสมาชิก (Register System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: กรอกข้อมูล username, email, password", ""),
            ("->", "2: POST /api/auth/register\n(username, email, password)", ""),
            ("note", "ระบบตรวจสอบ username และ email\nไม่ซ้ำ, เข้ารหัส password ด้วย bcrypt\nสร้าง User + UserProfile", ""),
            ("<-", "3: return {message: 'Register success',\nuser_id}", ""),
            ("->", "4: กรอก email, password เพื่อเข้าสู่ระบบ", ""),
            ("->", "5: POST /api/auth/login\n(email, password)", ""),
            ("note", "ระบบตรวจสอบรหัสผ่าน bcrypt.verify()\nสร้าง JWT Token (หมดอายุ 1440 นาที)", ""),
            ("<-", "6: return {access_token, token_type: 'bearer',\nrole}", ""),
            ("note", "Frontend บันทึก token ไว้ใน localStorage\nผ่าน AuthContext", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การวิเคราะห์ใบหน้า (Face Analysis)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการวิเคราะห์ใบหน้าและเฉดสีผิว (Face Analysis System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: เลือกเพศ และอัปโหลดรูปภาพใบหน้า", ""),
            ("->", "2: POST /api/analysis/\n(gender, file: image)", ""),
            ("note", "Single-Pass Pipeline:\n• MediaPipe FaceMesh → landmarks 468 จุด\n• CNN ตรวจจับรูปทรงใบหน้า (7 shapes)\n• K-Means สกัดสีผิว → CIELAB (L*,a*,b*)\n• จำแนก skin_tone จาก L*\n• จำแนก undertone จาก a*, b*\n• CIEDE2000 จับคู่ Personal Color (12 sub-seasons)\n• สร้างรายการแนะนำสินค้า (S1–S6)", ""),
            ("<-", "3: return {analysis_id, face_shape,\nskin_tone, skin_undertone,\npersonal_color, palette,\nface_shape_tips, recommendations,\nrecommendations_by_category}", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การเรียกดูและค้นหาสินค้า (Browse Products)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการเรียกดูและค้นหาสินค้า (Browse Products System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: เข้าหน้าสินค้า", ""),
            ("->", "2: GET /api/products/categories", ""),
            ("<-", "3: return [หมวดหมู่สินค้าทั้งหมด]", ""),
            ("->", "4: GET /api/products/brands", ""),
            ("<-", "5: return [แบรนด์ที่มีสินค้า]", ""),
            ("->", "6: เลือกตัวกรอง (หมวดหมู่, แบรนด์, personal color,\nค้นหา, เรียงลำดับ)", ""),
            ("->", "7: GET /api/products/\n(?category_id, brand_id, personal_color,\nsearch, sort, skip, limit)", ""),
            ("note", "ระบบกรองสินค้าตามเงื่อนไข\nรองรับ pagination (skip/limit)\nเรียงลำดับ: ราคา, ใหม่สุด", ""),
            ("<-", "8: return [products: {product_id, name,\nprice, image_url, brand_name,\ncategory_name, links}]", ""),
            ("->", "9: คลิกดูรายละเอียดสินค้า", ""),
            ("->", "10: GET /api/products/{product_id}", ""),
            ("<-", "11: return {product details + color_shades\n+ links + reviews}", ""),
            ("->", "12: GET /api/products/similar/{product_id}", ""),
            ("note", "ระบบคำนวณความคล้ายคลึงด้วย\nTF-IDF + Cosine Similarity", ""),
            ("<-", "13: return {similar: [products],\nalgorithm: 'TF-IDF Cosine Similarity'}", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การแนะนำสินค้า (Recommendation)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการแนะนำสินค้า (Recommendation System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: ดูผลวิเคราะห์ / ขอรายการแนะนำ", ""),
            ("->", "2: GET /api/recommendations/{analysis_id}\n(?limit=20)", ""),
            ("note", "Multi-Signal Scoring Engine:\n• S1: Rule Match Priority (0–300)\n• S2: CIEDE2000 Color Proximity (0–300)\n• S3: Seasonal Match (0–150)\n• S4: Avoid-Color Penalty (-100–0)\n• S5: Feedback Aggregation (0–100)\n• S6: Popularity (0–50)\nเรียงลำดับตามคะแนนรวม", ""),
            ("<-", "3: return {recommendations: [{score,\nproduct}], by_category, total}", ""),
            ("->", "4: กดถูกใจ/ไม่ถูกใจ สินค้าที่แนะนำ", ""),
            ("->", "5: POST /api/recommendations/feedback/\n{recommendation_id} (rating: like/dislike)", ""),
            ("<-", "6: return {message: 'Feedback saved'}", ""),
            ("note", "Feedback ถูกนำไปคำนวณ S5\nในการแนะนำครั้งถัดไป", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การเปรียบเทียบสินค้า (Product Comparison)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการเปรียบเทียบสินค้า (Product Comparison System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: เลือกสินค้า 1–3 ชิ้นเพื่อเปรียบเทียบ", ""),
            ("->", "2: POST /api/recommendations/compare\n{analysis_id, product_ids: [1,2,3]}", ""),
            ("note", "ระบบตรวจสอบจำนวน product_ids (1–3)\nคำนวณ Color Match % สำหรับแต่ละสินค้า:\nmatch_pct = clamp((S2+S4)/300×100, 0, 100)\nดึงเคล็ดลับตามรูปหน้า", ""),
            ("<-", "3: return {face_shape, products: {\nid: {s2, s4, match_pct}}}", ""),
            ("note", "Frontend แสดงตารางเปรียบเทียบ\nป้ายกำกับภาษาไทย: ความเข้ากันของสี,\nเคล็ดลับตามรูปหน้า, ราคา, แบรนด์ ฯลฯ", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การสนทนากับ Gemini AI (AI Chat)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการสนทนากับ Gemini AI (AI Chat System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: เข้าหน้า Gemini Chat", ""),
            ("->", "2: POST /api/gemini/session\n(title: 'New Session')", ""),
            ("note", "สร้างเซสชันใหม่ เชื่อมโยงกับ\nAnalysisResult ล่าสุด (ถ้ามี)", ""),
            ("<-", "3: return {session_id, title}", ""),
            ("->", "4: พิมพ์คำถาม / อัปโหลดรูป", ""),
            ("->", "5: POST /api/gemini/session/{id}/chat\n(prompt, file: image [optional])", ""),
            ("note", "ระบบรวบรวมประวัติจาก GeminiMessage\nสร้าง context จากผลวิเคราะห์ (ถ้ามี)\nส่งไปยัง Google Generative AI API", ""),
            ("<-", "6: return {message_id, prompt,\nresponse, image_output}", ""),
            ("->", "7: ดูประวัติการสนทนา", ""),
            ("->", "8: GET /api/gemini/session/{id}/messages", ""),
            ("<-", "9: return [GeminiMessage[]]", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การแต่งรูปและบันทึกลุค (Photo Editor & Saved Looks)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการแต่งรูปและบันทึกลุค (Photo Editor System Sequence Diagram)",
        "actor": "ผู้ใช้งาน\n(User)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: เข้าหน้า Editor / เลือกรูปภาพ", ""),
            ("->", "2: GET /api/products/makeup-recommendations\n(?parts=lips,eyes,blush&limit=4)", ""),
            ("<-", "3: return {lips: [products],\neyes: [products], blush: [products]}", ""),
            ("->", "4: เลือกเครื่องสำอาง / ฟิลเตอร์ / สติกเกอร์\nปรับค่าความสว่าง คอนทราสต์", ""),
            ("->", "5: POST /api/looks/\n{name, category, makeup_data, filter_data}", ""),
            ("note", "บันทึกลุคเครื่องสำอางพร้อม\nค่าฟิลเตอร์ทั้งหมดในรูปแบบ JSON", ""),
            ("<-", "6: return {look_id, name,\nmessage: 'บันทึกลุคสำเร็จ'}", ""),
            ("->", "7: ดูลุคที่บันทึกไว้", ""),
            ("->", "8: GET /api/looks/", ""),
            ("<-", "9: return [SavedLook[]]", ""),
            ("->", "10: ลบลุค", ""),
            ("->", "11: DELETE /api/looks/{look_id}", ""),
            ("<-", "12: return {message: 'ลบลุคสำเร็จ'}", ""),
        ],
    },
    {
        "title": "System Sequence Diagram: การติดตามพฤติกรรมผู้ใช้ (Behavior Tracking & Analytics)",
        "caption": "ภาพที่ 4.x แผนภาพลำดับระบบการติดตามพฤติกรรมผู้ใช้ (Behavior Tracking System Sequence Diagram)",
        "actor": "ผู้ใช้งาน / แอดมิน\n(User / Admin)",
        "system": "ระบบ AuraMatch\n(System)",
        "messages": [
            ("->", "1: ใช้งานเว็บไซต์ (ดูสินค้า, ค้นหา, กรอง, คลิก)", ""),
            ("->", "2: POST /api/behavior/track\n{event_type, event_data, page, session_id}", ""),
            ("note", "event_type: product_view, search,\nfilter, click, similar_view,\nmakeup_select, preset_apply", ""),
            ("<-", "3: return {status: 'tracked'}", ""),
            ("->", "4: [Admin] ดูสรุปสถิติการใช้งาน", ""),
            ("->", "5: GET /api/behavior/analytics/summary\n(?days=30)", ""),
            ("<-", "6: return {total_events, unique_users,\nunique_sessions, event_breakdown}", ""),
            ("->", "7: GET /api/behavior/analytics/top-products\n(?days=30&limit=10)", ""),
            ("<-", "8: return [{product_id, product_name,\nviews}]", ""),
            ("->", "9: GET /api/behavior/analytics/click-funnel\n(?days=30)", ""),
            ("<-", "10: return {product_views,\npurchase_clicks, view_to_click_rate}", ""),
        ],
    },
]


# ═══════════════════════════════════════════
# Drawing Engine
# ═══════════════════════════════════════════

def _hex(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16)/255 for i in (0,2,4))


def render_ssd(ssd):
    actor_label = ssd["actor"]
    system_label = ssd["system"]
    messages = ssd["messages"]

    # Layout constants
    actor_x = 3.0
    system_x = 18.0
    box_w = 5.0
    box_h = 1.0
    start_y = 0
    step_y = 1.6   # vertical spacing per message
    note_w = 8.5

    # Count total vertical steps (notes take 1 step too)
    total_steps = len(messages)
    fig_h = max((total_steps + 3) * step_y, 8)
    fig_w = 24

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(-1, fig_w + 1)
    ax.set_ylim(-fig_h - 1, 2.5)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # ─── Actor box (top) ───
    actor_box = FancyBboxPatch(
        (actor_x - box_w/2, start_y - box_h), box_w, box_h,
        boxstyle="round,pad=0.1",
        facecolor=_hex(ACTOR_HEAD), edgecolor=_hex(ARROW_C), linewidth=1.8, zorder=3)
    ax.add_patch(actor_box)
    ax.text(actor_x, start_y - box_h/2, actor_label,
            ha="center", va="center", fontsize=HEAD_FS, fontweight="bold",
            fontfamily=FONT, zorder=4)

    # ─── System box (top) ───
    sys_box = FancyBboxPatch(
        (system_x - box_w/2, start_y - box_h), box_w, box_h,
        boxstyle="round,pad=0.1",
        facecolor=_hex(SYSTEM_HEAD), edgecolor=_hex(ARROW_C), linewidth=1.8, zorder=3)
    ax.add_patch(sys_box)
    ax.text(system_x, start_y - box_h/2, system_label,
            ha="center", va="center", fontsize=HEAD_FS, fontweight="bold",
            fontfamily=FONT, zorder=4)

    # ─── Lifelines ───
    lifeline_end_y = -(total_steps + 1.5) * step_y
    ax.plot([actor_x, actor_x], [start_y - box_h, lifeline_end_y],
            color=_hex(LIFELINE_C), linewidth=1.2, linestyle="--", zorder=1)
    ax.plot([system_x, system_x], [start_y - box_h, lifeline_end_y],
            color=_hex(LIFELINE_C), linewidth=1.2, linestyle="--", zorder=1)

    # ─── Actor box (bottom) ───
    actor_bot = FancyBboxPatch(
        (actor_x - box_w/2, lifeline_end_y - box_h), box_w, box_h,
        boxstyle="round,pad=0.1",
        facecolor=_hex(ACTOR_HEAD), edgecolor=_hex(ARROW_C), linewidth=1.8, zorder=3)
    ax.add_patch(actor_bot)
    ax.text(actor_x, lifeline_end_y - box_h/2, actor_label,
            ha="center", va="center", fontsize=HEAD_FS, fontweight="bold",
            fontfamily=FONT, zorder=4)

    sys_bot = FancyBboxPatch(
        (system_x - box_w/2, lifeline_end_y - box_h), box_w, box_h,
        boxstyle="round,pad=0.1",
        facecolor=_hex(SYSTEM_HEAD), edgecolor=_hex(ARROW_C), linewidth=1.8, zorder=3)
    ax.add_patch(sys_bot)
    ax.text(system_x, lifeline_end_y - box_h/2, system_label,
            ha="center", va="center", fontsize=HEAD_FS, fontweight="bold",
            fontfamily=FONT, zorder=4)

    # ─── Messages ───
    y = start_y - box_h - step_y
    for direction, label, _ in messages:
        if direction == "->":
            # Actor → System (solid arrow)
            ax.annotate("",
                xy=(system_x - 0.3, y), xytext=(actor_x + 0.3, y),
                arrowprops=dict(arrowstyle="-|>", color=_hex(ARROW_C),
                                linewidth=1.8, mutation_scale=18),
                zorder=5)
            mid_x = (actor_x + system_x) / 2
            n_lines = label.count("\n") + 1
            ax.text(mid_x, y + 0.25 + (n_lines - 1) * 0.15, label,
                    ha="center", va="bottom", fontsize=MSG_FS,
                    fontfamily=FONT, color=_hex(ARROW_C), zorder=5)

        elif direction == "<-":
            # System → Actor (dashed return arrow)
            ax.annotate("",
                xy=(actor_x + 0.3, y), xytext=(system_x - 0.3, y),
                arrowprops=dict(arrowstyle="-|>", color=_hex(RETURN_C),
                                linewidth=1.8, linestyle="--", mutation_scale=18),
                zorder=5)
            mid_x = (actor_x + system_x) / 2
            n_lines = label.count("\n") + 1
            ax.text(mid_x, y + 0.25 + (n_lines - 1) * 0.15, label,
                    ha="center", va="bottom", fontsize=MSG_FS,
                    fontfamily=FONT, color=_hex(RETURN_C), zorder=5)

        elif direction == "note":
            # Note box on system side
            note_lines = label.split("\n")
            note_h = len(note_lines) * 0.38 + 0.3
            note_x = system_x + 1.0
            note_rect = FancyBboxPatch(
                (note_x, y - note_h/2), note_w, note_h,
                boxstyle="round,pad=0.15",
                facecolor=_hex(NOTE_BG), edgecolor=_hex(NOTE_BORDER),
                linewidth=1.0, zorder=3)
            ax.add_patch(note_rect)
            # Connector line
            ax.plot([system_x + 0.1, note_x], [y, y],
                    color=_hex(NOTE_BORDER), linewidth=0.8, linestyle=":", zorder=2)
            text_y = y + note_h/2 - 0.35
            for nl in note_lines:
                ax.text(note_x + 0.25, text_y, nl, ha="left", va="center",
                        fontsize=NOTE_FS, fontfamily=FONT, color="#424242", zorder=4)
                text_y -= 0.38

        y -= step_y

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════
# Build Word Document
# ═══════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ── Cover ──
for _ in range(3):
    doc.add_paragraph("")

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AuraMatch")
r.font.size = Pt(40); r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("System Sequence Diagrams")
r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph("")

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("แผนภาพลำดับระบบ (System Sequence Diagram)\nแสดงลำดับการสื่อสารระหว่าง Actor กับระบบในมุมมอง Black-Box")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_paragraph("")

i = doc.add_paragraph()
i.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = i.add_run(f"{len(SSDS)} Use Cases  ·  ภาพรวมระบบ AuraMatch ทั้งหมด")
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

doc.add_paragraph("")

# Legend
leg = doc.add_paragraph()
leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = leg.add_run("สัญลักษณ์ (Legend)")
r.font.size = Pt(14); r.font.bold = True

legends = [
    "── ลูกศรทึบสีดำ →  = ข้อความจาก Actor ไปยัง System (Request)",
    "- - ลูกศรประสีน้ำเงิน ← = ข้อความตอบกลับจาก System (Response)",
    "| กล่องสีครีม = หมายเหตุ (Note) อธิบายการประมวลผลภายในระบบ",
    "| กล่องสีฟ้า = Actor (ผู้ใช้งาน)  |  กล่องสีเขียว = System (ระบบ)",
    "- - เส้นประแนวตั้ง = Lifeline (เส้นชีวิตของออบเจกต์)",
]
for item in legends:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

# ── TOC ──
doc.add_page_break()
toc = doc.add_paragraph()
toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc.add_run("สารบัญ — Table of Contents")
r.font.size = Pt(20); r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph("")

for idx, ssd in enumerate(SSDS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    r = p.add_run(f"หน้า {idx + 2}  —  {ssd['title']}")
    r.font.size = Pt(12); r.font.bold = True

# ── Diagram pages ──
for ssd in SSDS:
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ssd["title"])
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    img = render_ssd(ssd)
    doc.add_picture(img, width=Cm(25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(ssd["caption"])
    r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

# ── Save ──
out = os.path.join(OUT_DIR, "AuraMatch_System_Sequence_Diagrams.docx")
doc.save(out)
print(f"Saved: {out}")
